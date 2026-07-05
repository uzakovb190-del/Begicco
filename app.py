import streamlit as st
from supabase import create_client, Client
from datetime import date
import base64
import hmac
import os
import mimetypes

# ============================================
# CONFIG
# ============================================
# Credentials are read from Streamlit secrets, NOT hardcoded.
# - Local dev:  put them in  .streamlit/secrets.toml  (git-ignored)
# - Deployed:   Streamlit Cloud → your app → Settings → Secrets
_secrets_found = True
_available_keys = []
try:
    _available_keys = list(st.secrets.keys())
except Exception:
    _secrets_found = False

if not _secrets_found:
    st.error(
        "No Streamlit secrets were found at all.\n\n"
        "• Locally: the file must be at  .streamlit/secrets.toml  in the SAME folder "
        "you run `streamlit run app.py` from (not inside another subfolder), and named "
        "exactly secrets.toml (not secrets.toml.txt).\n"
        "• On Streamlit Cloud: open your app → Settings → Secrets, paste the values, Save, "
        "and let it reboot."
    )
    st.stop()

_missing = [k for k in ("SUPABASE_URL", "SUPABASE_KEY") if k not in st.secrets]
if _missing:
    st.error(
        f"Secrets were found, but these required keys are missing: {_missing}\n\n"
        f"Keys I can actually see right now: {_available_keys}\n\n"
        "Fix: in your secrets, the lines must be exactly (note the spelling/caps, with quotes, "
        "and NOT under any [section] header):\n"
        'SUPABASE_URL = "https://...supabase.co"\n'
        'SUPABASE_KEY = "eyJ..."'
    )
    st.stop()

SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

# ---------------------------------------------------------------------------
# OPTIONAL LOCAL BACKGROUND IMAGE
# To add a background: drop an image file in your repo (e.g. create an
# "assets" folder and put "background.jpg" inside it), then set the path here.
# Leave it as None to keep using the wallpaper uploaded from the Settings page.
# ---------------------------------------------------------------------------
LOCAL_WALLPAPER_PATH = None  # e.g. "assets/background.jpg"

@st.cache_resource
def init_supabase() -> Client:
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase = init_supabase()

st.set_page_config(
    page_title="Life Archive",
    page_icon="🗂️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============================================
# LOAD SETTINGS (cached — no more re-fetching on every rerun)
# ============================================
WALLPAPER_BUCKET = "wallpapers"

@st.cache_data(ttl=3600, show_spinner=False)
def get_setting(key: str):
    """Read one value from the settings table. Cached for an hour;
    call get_setting.clear() after any write to refresh immediately."""
    try:
        r = supabase.table("settings").select("value").eq("key", key).execute()
        if r.data:
            return r.data[0]["value"]
    except Exception:
        pass  # settings are cosmetic — fall back to defaults silently
    return None

def get_overlay_opacity() -> float:
    v = get_setting("overlay_opacity")
    try:
        return float(v) if v is not None else 0.6
    except (TypeError, ValueError):
        return 0.6

def _local_wallpaper():
    """Return a CSS-ready data URI for the local wallpaper file, or None."""
    if LOCAL_WALLPAPER_PATH and os.path.exists(LOCAL_WALLPAPER_PATH):
        try:
            with open(LOCAL_WALLPAPER_PATH, "rb") as f:
                data = f.read()
            mime = mimetypes.guess_type(LOCAL_WALLPAPER_PATH)[0] or "image/jpeg"
            return f"data:{mime};base64,{base64.b64encode(data).decode('utf-8')}"
        except Exception:
            return None
    return None

# Resolution order: local file → Storage URL (new, fast) → legacy base64 (old rows)
wallpaper_src = _local_wallpaper()
if not wallpaper_src:
    _wp_url = get_setting("wallpaper_url")
    if _wp_url:
        wallpaper_src = _wp_url
    else:
        _wp_b64 = get_setting("wallpaper_b64")  # backward-compat with old saves
        if _wp_b64:
            wallpaper_src = f"data:image/jpeg;base64,{_wp_b64}"
overlay_opacity = get_overlay_opacity()

# ============================================
# GLOBAL STYLES — SOFT NOIR THEME
# ============================================
wallpaper_css = ""
if wallpaper_src:
    wallpaper_css = f"""
    .stApp {{
        background-image: url("{wallpaper_src}");
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
    }}
    .stApp::before {{
        content: '';
        position: fixed;
        top: 0; left: 0; right: 0; bottom: 0;
        background: rgba(10, 10, 26, {overlay_opacity});
        z-index: 0;
        pointer-events: none;
    }}
    """

st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;700&family=Syne:wght@400;700;800&family=Playfair+Display:ital,wght@0,400;0,700;1,400;1,700&display=swap');

    {wallpaper_css}

    html, body, [class*="css"] {{
        font-family: 'Syne', sans-serif;
        background-color: #0a0a1a;
        color: #e8e0ff;
    }}

    section[data-testid="stSidebar"] {{
        background-color: #0d0d22 !important;
        border-right: 1px solid #2a2a4a !important;
    }}

    .main .block-container {{
        padding-top: 2rem;
        padding-bottom: 2rem;
    }}

    .section-header {{
        font-family: 'Syne', sans-serif;
        font-size: 1.8rem;
        font-weight: 800;
        letter-spacing: -0.5px;
        margin-bottom: 0.2rem;
        color: #e8e0ff;
    }}

    .section-sub {{
        font-size: 0.85rem;
        color: #3a3a6a;
        margin-bottom: 1.5rem;
        font-family: 'JetBrains Mono', monospace;
    }}

    .badge {{
        display: inline-block;
        padding: 2px 10px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 700;
        font-family: 'JetBrains Mono', monospace;
        letter-spacing: 0.5px;
    }}
    .badge-green  {{ background: #1a3a2a; color: #4ade80; border: 1px solid #4ade80; }}
    .badge-yellow {{ background: #3a2e0a; color: #facc15; border: 1px solid #facc15; }}
    .badge-red    {{ background: #3a1212; color: #f87171; border: 1px solid #f87171; }}
    .badge-blue   {{ background: #1a1a3a; color: #818cf8; border: 1px solid #818cf8; }}
    .badge-orange {{ background: #3a1f0a; color: #fb923c; border: 1px solid #fb923c; }}
    .badge-grey   {{ background: #1e1e2e; color: #9ca3af; border: 1px solid #444; }}
    .badge-gold   {{ background: #3a2e00; color: #fbbf24; border: 1px solid #fbbf24; }}
    .badge-purple {{ background: #1a1a3a; color: #a78bfa; border: 1px solid #a78bfa; }}

    .card {{
        background: #0d0d22;
        border: 1px solid #2a2a4a;
        border-radius: 12px;
        padding: 1.2rem 1.4rem;
        margin-bottom: 1rem;
    }}

    .divider {{
        border: none;
        border-top: 1px solid #2a2a4a;
        margin: 1.5rem 0;
    }}

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div {{
        background-color: #0d0d22 !important;
        border: 1px solid #2a2a4a !important;
        color: #e8e0ff !important;
        border-radius: 8px !important;
    }}

    .stButton > button {{
        background-color: #0d0d22;
        color: #e8e0ff;
        border: 1px solid #2a2a4a;
        border-radius: 8px;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.85rem;
        padding: 0.4rem 1.2rem;
        transition: all 0.2s ease;
    }}
    .stButton > button:hover {{
        border-color: #a78bfa;
        color: #a78bfa;
        background-color: #1a1a3a;
    }}

    .stButton > button[kind="primary"] {{
        background-color: #1a1a3a;
        border-color: #a78bfa;
        color: #a78bfa;
    }}

    .metric-row {{
        display: flex;
        gap: 1rem;
        margin-bottom: 1.5rem;
    }}
    .metric-card {{
        flex: 1;
        background: #0d0d22;
        border: 1px solid #2a2a4a;
        border-radius: 10px;
        padding: 1rem;
        text-align: center;
        transition: all 0.2s ease;
    }}
    .metric-card:hover {{
        border-color: #a78bfa;
        background: #1a1a3a;
    }}
    .metric-value {{
        font-size: 2rem;
        font-weight: 800;
        font-family: 'JetBrains Mono', monospace;
    }}
    .metric-label {{
        font-size: 0.75rem;
        color: #3a3a6a;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-top: 4px;
    }}

    .chivalry {{
        font-family: 'Playfair Display', serif;
        font-style: italic;
        color: #9a90cc;
    }}

    /* Page transition */
    @keyframes fadeInPage {{
        from {{ opacity: 0; transform: translateY(10px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    .main .block-container {{
        animation: fadeInPage 0.4s ease forwards;
    }}

    /* Progress bar */
    .progress-bar-bg {{
        background: #1a1a3a;
        border-radius: 10px;
        height: 8px;
        width: 100%;
        margin: 6px 0;
    }}
    .progress-bar-fill {{
        background: linear-gradient(90deg, #a78bfa, #818cf8);
        border-radius: 10px;
        height: 8px;
        transition: width 0.3s ease;
    }}
    /* Sidebar toggle buttons — keep BOTH the collapse button (sidebar OPEN)
       and the expand button (sidebar COLLAPSED) visible and styled.
       NOTE: in Streamlit, the expand button lives INSIDE the toolbar/header,
       so we must NOT hide stToolbar or the header — only the deploy/menu items. */
    [data-testid="stSidebarCollapseButton"],
    [data-testid="stExpandSidebarButton"],
    [data-testid="collapsedControl"] {{
        background-color: #1a1a3a !important;
        border: 1px solid #a78bfa !important;
        border-radius: 8px !important;
        opacity: 1 !important;
        visibility: visible !important;
    }}
    [data-testid="stSidebarCollapseButton"] svg,
    [data-testid="stExpandSidebarButton"] svg,
    [data-testid="collapsedControl"] svg {{
        fill: #a78bfa !important;
    }}
    /* Keep the header/toolbar present (so the expand button survives), just
       transparent, and hide only the deploy button / hamburger menu / status. */
    header[data-testid="stHeader"] {{
        background: transparent !important;
    }}
    [data-testid="stToolbarActions"] {{display: none !important;}}
    [data-testid="stMainMenu"] {{display: none !important;}}
    [data-testid="stAppDeployButton"] {{display: none !important;}}
    [data-testid="stStatusWidget"] {{display: none !important;}}
    [data-testid="stDecoration"] {{display: none !important;}}
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
</style>
""", unsafe_allow_html=True)

# ============================================
# BADGE HELPERS
# ============================================
def mood_badge(score: int) -> str:
    if score >= 8:
        return f'<span class="badge badge-green">😊 {score}/10</span>'
    elif score >= 5:
        return f'<span class="badge badge-yellow">😐 {score}/10</span>'
    else:
        return f'<span class="badge badge-red">😔 {score}/10</span>'

def clarity_badge(clarity: str) -> str:
    colors = {"sharp": "badge-green", "normal": "badge-blue", "foggy": "badge-orange"}
    icons  = {"sharp": "⚡", "normal": "🔵", "foggy": "🌫️"}
    return f'<span class="badge {colors.get(clarity,"badge-grey")}">{icons.get(clarity,"●")} {clarity}</span>'

def physical_badge(state: str) -> str:
    colors = {"energized": "badge-green", "neutral": "badge-blue", "tired": "badge-red"}
    icons  = {"energized": "⚡", "neutral": "🔵", "tired": "🔴"}
    return f'<span class="badge {colors.get(state,"badge-grey")}">{icons.get(state,"●")} {state}</span>'

def mental_badge(state: str) -> str:
    colors = {"calm": "badge-green", "stable": "badge-blue", "stressed": "badge-orange", "heavy": "badge-red"}
    icons  = {"calm": "🟢", "stable": "🔵", "stressed": "🟠", "heavy": "🔴"}
    return f'<span class="badge {colors.get(state,"badge-grey")}">{icons.get(state,"●")} {state}</span>'

def result_badge(result: str) -> str:
    colors = {"win": "badge-green", "pass": "badge-blue", "fail": "badge-red", "complete": "badge-gold"}
    icons  = {"win": "🏆", "pass": "✅", "fail": "❌", "complete": "⭐"}
    return f'<span class="badge {colors.get(result,"badge-grey")}">{icons.get(result,"●")} {result.upper()}</span>'

def goal_status_badge(status: str) -> str:
    colors = {"active": "badge-green", "paused": "badge-grey", "completed": "badge-gold"}
    icons  = {"active": "🟢", "paused": "⏸️", "completed": "⭐"}
    return f'<span class="badge {colors.get(status,"badge-grey")}">{icons.get(status,"●")} {status.upper()}</span>'

def significance_badge(score: int) -> str:
    if score >= 5:
        return f'<span class="badge badge-gold">★★★★★ {score}/5</span>'
    elif score >= 4:
        return f'<span class="badge badge-orange">★★★★☆ {score}/5</span>'
    elif score >= 3:
        return f'<span class="badge badge-yellow">★★★☆☆ {score}/5</span>'
    else:
        return f'<span class="badge badge-grey">★★☆☆☆ {score}/5</span>'

def skill_level_badge(level: str) -> str:
    colors = {"beginner": "badge-blue", "intermediate": "badge-yellow", "advanced": "badge-orange", "master": "badge-gold"}
    return f'<span class="badge {colors.get(level,"badge-grey")}">⚡ {level.upper()}</span>'

# ============================================
# LOGIN GATE
# ============================================
# Set APP_PASSCODE in your Streamlit secrets to require a code before entry.
# If APP_PASSCODE is not set, the app stays open (so you can never lock yourself
# out by forgetting to configure it).
APP_PASSCODE = st.secrets.get("APP_PASSCODE", None)

if APP_PASSCODE and not st.session_state.get("authed", False):
    st.markdown("<div style='height:8vh'></div>", unsafe_allow_html=True)
    _lc1, _lc2, _lc3 = st.columns([1, 1.4, 1])
    with _lc2:
        st.markdown("""
        <div class="card" style="text-align:center;">
            <div style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:800;color:#e8e0ff;">🗂️ Life Archive</div>
            <div style="font-family:'JetBrains Mono',monospace;font-size:0.75rem;color:#3a3a6a;margin-top:6px;">enter your access code to continue</div>
        </div>
        """, unsafe_allow_html=True)
        with st.form("login_form"):
            _code = st.text_input("Access code", type="password",
                                  label_visibility="collapsed", placeholder="Access code")
            _entered = st.form_submit_button("Enter →", use_container_width=True, type="primary")
        if _entered:
            if hmac.compare_digest(str(_code), str(APP_PASSCODE)):
                st.session_state["authed"] = True
                st.rerun()
            else:
                st.error("Incorrect code.")
    st.stop()

# ============================================
# NAVIGATION
# ============================================
nav_options = [
    "🏠  Home",
    "📝  Daily Log",
    "📖  Reading Log",
    "🚨  Life Event",
    "🛍️  Purchase Tracker",
    "💫  Wish List",
    "🎯  Goals",
    "🏆  Outcomes",
    "🧠  Skills",
    "📜  Archive",
    "⚙️  Settings",
]

# ---- In-page navigation (no sidebar) ----
if "nav" not in st.session_state:
    st.session_state["nav"] = "🏠  Home"
if "current_page_override" not in st.session_state:
    st.session_state["current_page_override"] = None

# Bridge: existing buttons set "current_page_override" — honor it as a nav change.
if st.session_state.get("current_page_override"):
    if st.session_state["current_page_override"] in nav_options:
        st.session_state["nav"] = st.session_state["current_page_override"]
    st.session_state["current_page_override"] = None

def go_to(dest):
    """Navigate to a section (one layer in) or back Home."""
    st.session_state["nav"] = dest
    st.rerun()

page = st.session_state["nav"]

# ---- Top bar: brand on the left, a Back-to-Home button on every non-Home page ----
_tl, _tr = st.columns([4, 1])
with _tl:
    st.markdown("""
        <div style="display:flex;align-items:baseline;gap:0.6rem;flex-wrap:wrap;">
            <span style="font-family:'Syne',sans-serif;font-size:1.15rem;font-weight:800;color:#e8e0ff;">🗂️ Life Archive</span>
            <span style="font-family:'JetBrains Mono',monospace;font-size:0.65rem;color:#3a3a6a;">personal memory system</span>
        </div>
    """, unsafe_allow_html=True)
with _tr:
    if page != "🏠  Home":
        if st.button("← Home", key="nav_back_home", use_container_width=True):
            go_to("🏠  Home")
st.markdown('<hr class="divider" style="margin-top:0.5rem;">', unsafe_allow_html=True)

# ============================================
# HOME PAGE
# ============================================
if page == "🏠  Home":
    st.markdown("<h1 style='font-family:Syne,sans-serif;font-size:2.2rem;font-weight:800;color:#e8e0ff;'>Welcome back, Captain 🫡</h1>", unsafe_allow_html=True)
    st.markdown(f'<div class="section-sub">Today is {date.today().strftime("%A, %B %d %Y")} · Your archive is running.</div>', unsafe_allow_html=True)

    try:
        total_logs     = supabase.table("daily_logs").select("id", count="exact").execute().count or 0
        total_goals    = supabase.table("goals").select("id", count="exact").execute().count or 0
        total_wishes   = supabase.table("wishes").select("id", count="exact").execute().count or 0
        total_outcomes = supabase.table("outcomes").select("id", count="exact").execute().count or 0
    except Exception:
        total_logs = total_goals = total_wishes = total_outcomes = 0

    # Clickable metric cards — the number itself is the button.
    st.markdown("""
    <style>
    .st-key-m_logs button, .st-key-m_goals button,
    .st-key-m_wishes button, .st-key-m_outcomes button {
        height: 118px;
        background: #0d0d22 !important;
        border: 1px solid #2a2a4a !important;
        border-radius: 12px !important;
        transition: all 0.2s ease;
    }
    .st-key-m_logs button:hover, .st-key-m_goals button:hover,
    .st-key-m_wishes button:hover, .st-key-m_outcomes button:hover {
        border-color: #a78bfa !important;
        background: #1a1a3a !important;
    }
    .st-key-m_logs button p, .st-key-m_goals button p,
    .st-key-m_wishes button p, .st-key-m_outcomes button p {
        font-size: 2.4rem !important;
        font-weight: 800 !important;
        font-family: 'JetBrains Mono', monospace !important;
    }
    .st-key-m_logs button p, .st-key-m_logs button { color: #4ade80 !important; }
    .st-key-m_goals button p, .st-key-m_goals button { color: #a78bfa !important; }
    .st-key-m_wishes button p, .st-key-m_wishes button { color: #c084fc !important; }
    .st-key-m_outcomes button p, .st-key-m_outcomes button { color: #818cf8 !important; }
    </style>
    """, unsafe_allow_html=True)

    def metric_link(col, value, label, key, dest, sub=None):
        with col:
            if st.button(str(value), key=key, use_container_width=True):
                if sub:
                    st.session_state["archive_view"] = sub
                go_to(dest)
            st.markdown(
                f'<div class="metric-label" style="text-align:center;margin-top:-6px;">{label}</div>',
                unsafe_allow_html=True,
            )

    col1, col2, col3, col4 = st.columns(4)
    metric_link(col1, total_logs,     "Days Logged",  "m_logs",     "📜  Archive", sub="📅 Daily Log History")
    metric_link(col2, total_goals,    "Active Goals", "m_goals",    "🎯  Goals")
    metric_link(col3, total_wishes,   "Wishes",       "m_wishes",   "💫  Wish List")
    metric_link(col4, total_outcomes, "Outcomes",     "m_outcomes", "🏆  Outcomes")

    # ---- Navigation hub: tap a section to go one layer in ----
    st.markdown('<div class="section-sub" style="margin-top:1.4rem;letter-spacing:1px;">OPEN A SECTION</div>', unsafe_allow_html=True)
    _sections = [s for s in nav_options if s != "🏠  Home"]
    _per_row = 4
    for _i in range(0, len(_sections), _per_row):
        _row = _sections[_i:_i + _per_row]
        _cols = st.columns(_per_row)
        for _c, _sec in zip(_cols, _row):
            with _c:
                if st.button(_sec, key=f"hub_{_sec}", use_container_width=True):
                    go_to(_sec)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ---- Stats Snapshot Share ----
    with st.expander("📤 Share your stats snapshot"):
        snapshot_lines = [
            f"📊 *Life Archive — Stats Snapshot*",
            f"📅 {date.today().strftime('%B %d, %Y')}",
            f"",
            f"📝 Days Logged: {total_logs}",
            f"🎯 Active Goals: {total_goals}",
            f"💫 Wishes: {total_wishes}",
            f"🏆 Outcomes: {total_outcomes}",
            f"",
            f"— sent from Life Archive",
        ]
        snapshot_text = "\n".join(snapshot_lines)
        st.text_area("Copy and send to Telegram or your channel",
                     value=snapshot_text, height=200, key="home_share_box")
        tg_snap_url = f"https://t.me/share/url?url=&text={snapshot_text.replace(' ','%20').replace(chr(10),'%0A')}"
        _snap_dl, _snap_tg = st.columns([1, 2])
        with _snap_dl:
            import json as _json
            st.download_button(
                label="⬇️ Download snapshot",
                data=_json.dumps({
                    "date": str(date.today()),
                    "days_logged": total_logs,
                    "active_goals": total_goals,
                    "wishes": total_wishes,
                    "outcomes": total_outcomes,
                }, indent=2),
                file_name=f"life_archive_snapshot_{date.today()}.json",
                mime="application/json",
                use_container_width=True,
            )
        with _snap_tg:
            st.markdown(f'<a href="{tg_snap_url}" target="_blank" style="display:inline-block;width:100%;margin-top:0.3rem;padding:0.4rem 1.2rem;background:#2a2a4a;border:1px solid #a78bfa;border-radius:8px;color:#e8e0ff;text-decoration:none;font-family:JetBrains Mono,monospace;font-size:0.8rem;text-align:center;">📨 Open in Telegram</a>', unsafe_allow_html=True)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    st.markdown("""
    <div class="card">
        <div style="font-size:0.75rem; color:#3a3a6a; font-family:'JetBrains Mono',monospace; margin-bottom:0.6rem; letter-spacing:1px;">WHAT THIS IS</div>
        <div class="chivalry" style="font-size:1.05rem; line-height:1.9;">
            This is your private life archive. Not a productivity tool. Not a habit tracker.<br>
            <em>A structured memory of who you are, what you tried, and what actually happened.</em><br><br>
            Every entry you make becomes a permanent, queryable record of your life.
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="card">
        <div style="font-size:0.75rem; color:#3a3a6a; font-family:'JetBrains Mono',monospace; margin-bottom:0.8rem; letter-spacing:1px;">QUICK START</div>
        <div class="chivalry" style="font-size:1rem; line-height:2.2;">
            📝 &nbsp;<b>Daily Log</b> — start here every day<br>
            💫 &nbsp;<b>Wish List</b> — add anything you want to pursue<br>
            🎯 &nbsp;<b>Goals</b> — activate a wish and start tracking<br>
            🏆 &nbsp;<b>Outcomes</b> — record what happened when it ended<br>
            🧠 &nbsp;<b>Skills</b> — track mastery hour by hour<br>
            📜 &nbsp;<b>Archive</b> — browse your full history
        </div>
    </div>
    """, unsafe_allow_html=True)

elif page == "📝  Daily Log":
    import streamlit.components.v1 as components

    st.markdown('<div class="section-header">📝 Daily Log</div>', unsafe_allow_html=True)
    today = date.today()
    st.markdown(f'<div class="section-sub">{today.strftime("%A, %B %d %Y")} · morning pages when you wake · evening reflection before sleep</div>', unsafe_allow_html=True)

    # ------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------
    def _load_day(d):
        try:
            r = supabase.table("daily_logs").select("*").eq("date", str(d)).execute()
            return r.data[0] if r.data else None
        except Exception as ex:
            st.caption(f"⚠ couldn't load today's entry: {ex}")
            return None

    def _save_partial(fields: dict):
        """Insert or update only the given columns for today — so saving the
        morning never wipes the evening, and vice versa."""
        try:
            existing_row = _load_day(today)
            if existing_row:
                supabase.table("daily_logs").update(fields).eq("date", str(today)).execute()
            else:
                supabase.table("daily_logs").insert({"date": str(today), **fields}).execute()
            _diary_stats.clear()
            return True, None
        except Exception as ex:
            return False, str(ex)

    @st.cache_data(ttl=300, show_spinner=False)
    def _diary_stats():
        """Writing streak + lifetime word count — the 'watch it grow' numbers."""
        try:
            rows = supabase.table("daily_logs").select(
                "date,morning_entry,evening_entry,daily_summary,self_assessment"
            ).order("date", desc=True).limit(1000).execute().data or []
        except Exception:
            return 0, 0, 0
        total_words = 0
        for r in rows:
            for f in ("morning_entry", "evening_entry", "daily_summary", "self_assessment"):
                if r.get(f):
                    total_words += len(str(r[f]).split())
        # streak: consecutive days ending today or yesterday
        from datetime import timedelta
        logged = {r["date"] for r in rows}
        streak = 0
        cursor = today if str(today) in logged else today - timedelta(days=1)
        while str(cursor) in logged:
            streak += 1
            cursor -= timedelta(days=1)
        return streak, total_words, len(rows)

    @st.cache_data(ttl=3600, show_spinner=False)
    def _on_this_day():
        """Entries from 7 / 30 / 365 days ago — past-you says hi."""
        from datetime import timedelta
        out = []
        for days, label in [(7, "1 week ago"), (30, "1 month ago"), (365, "1 year ago")]:
            d = today - timedelta(days=days)
            try:
                r = supabase.table("daily_logs").select(
                    "date,evening_entry,morning_entry,daily_summary,mood_score,gratitude"
                ).eq("date", str(d)).execute()
                if r.data:
                    out.append((label, r.data[0]))
            except Exception:
                pass
        return out

    EVENING_PROMPTS = [
        "What did today teach you that yesterday-you didn't know?",
        "Where did you act like the person you want to become?",
        "What was the hardest moment today — and how did you handle it?",
        "What would you tell your younger brother about a day like this?",
        "What are you avoiding right now? Say it plainly.",
        "One thing you did today that future-you will thank you for.",
        "If today repeated 100 times, what one change would matter most?",
        "What drained you today? What refilled you?",
        "Describe today in one honest sentence. Then explain it.",
        "What small win did you almost not notice?",
        "Who helped you today — even in a tiny way?",
        "What did you do today purely because it was right?",
        "Kyoto is far. What did you do today that moved you 1 cm closer?",
        "What are you proud of that nobody else saw?",
        "What worried you this morning — and did it actually happen?",
        "What's one thing you want to remember about today in 10 years?",
    ]
    _prompt_today = EVENING_PROMPTS[today.toordinal() % len(EVENING_PROMPTS)]

    # ------------------------------------------------------------
    # Load state
    # ------------------------------------------------------------
    e = _load_day(today) or {}
    morning_done = bool(e.get("morning_entry") or e.get("morning_saved_at"))
    evening_done = bool(e.get("evening_entry") or e.get("evening_saved_at"))

    # per-day init of list widgets (prefill from saved row exactly once per day)
    if st.session_state.get("dl_loaded_date") != str(today):
        st.session_state["dl_loaded_date"] = str(today)
        st.session_state.accomplishments = e.get("accomplishments") or []
        st.session_state.media_list = e.get("media_consumed") or []
        st.session_state.workouts = e.get("workouts") or []

    # ------------------------------------------------------------
    # Growth header — streak · words · status
    # ------------------------------------------------------------
    _streak, _words, _days = _diary_stats()
    _m_badge = '<span class="badge badge-green">🌅 morning saved</span>' if morning_done else '<span class="badge badge-grey">🌅 morning pending</span>'
    _e_badge = '<span class="badge badge-green">🌙 evening saved</span>' if evening_done else '<span class="badge badge-grey">🌙 evening pending</span>'
    st.markdown(f"""
    <div class="card" style="display:flex;gap:1.6rem;flex-wrap:wrap;align-items:center;">
        <div><span style="font-family:'JetBrains Mono',monospace;font-size:1.5rem;font-weight:700;color:#facc15;">🔥 {_streak}</span>
             <div class="metric-label">day streak</div></div>
        <div><span style="font-family:'JetBrains Mono',monospace;font-size:1.5rem;font-weight:700;color:#a78bfa;">{_words:,}</span>
             <div class="metric-label">words written</div></div>
        <div><span style="font-family:'JetBrains Mono',monospace;font-size:1.5rem;font-weight:700;color:#4ade80;">{_days}</span>
             <div class="metric-label">days archived</div></div>
        <div style="margin-left:auto;">{_m_badge} &nbsp; {_e_badge}</div>
    </div>
    """, unsafe_allow_html=True)

    # ---- On this day ----
    _memories = _on_this_day()
    if _memories:
        with st.expander("🕰️ On this day — a note from past you"):
            for label, mem in _memories:
                snippet = (mem.get("evening_entry") or mem.get("morning_entry")
                           or mem.get("daily_summary") or "")
                snippet = snippet[:400] + ("…" if len(snippet) > 400 else "")
                grat = f"<br><span style='color:#4ade80;'>🙏 {mem['gratitude']}</span>" if mem.get("gratitude") else ""
                st.markdown(f"""
                <div class="card">
                    <div style="font-size:0.7rem;color:#3a3a6a;font-family:'JetBrains Mono',monospace;">{label} · {mem.get('date','')} · mood {mem.get('mood_score','—')}/10</div>
                    <div class="chivalry" style="margin-top:0.4rem;line-height:1.8;">{snippet}{grat}</div>
                </div>""", unsafe_allow_html=True)

    tab_morning, tab_evening, tab_metrics = st.tabs(["🌅 Morning Pages", "🌙 Evening Reflection", "📊 Day Metrics"])

    # ============================================================
    # 🌅 MORNING — write when you wake up
    # ============================================================
    with tab_morning:
        # a whisper from last night
        try:
            from datetime import timedelta
            _y = supabase.table("daily_logs").select("evening_entry,gratitude,intention") \
                .eq("date", str(today - timedelta(days=1))).execute().data
        except Exception:
            _y = None
        if _y and (_y[0].get("evening_entry") or _y[0].get("gratitude")):
            _last = (_y[0].get("evening_entry") or _y[0].get("gratitude"))[:220]
            st.markdown(f"""
            <div style="border-left:3px solid #a78bfa;padding:0.5rem 1rem;color:#9ca3af;font-style:italic;font-size:0.9rem;">
                last night you wrote: “{_last}…”
            </div><br>""", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            waking_mood = st.slider("Waking mood", 1, 10, int(e.get("waking_mood") or 6), key="m_waking_mood")
            sleep_duration = st.number_input("Sleep (hours)", 0.0, 24.0, float(e.get("sleep_duration") or 7.0), step=0.5, key="m_sleep")
        with col2:
            dream_log = st.text_area("Dream log (optional)", value=e.get("dream_log") or "", height=68,
                                     placeholder="Any dreams worth keeping?", key="m_dream")

        intention = st.text_input("🎯 One intention for today",
                                  value=e.get("intention") or "",
                                  placeholder="If only ONE thing happens today, it should be…", key="m_intention")

        morning_entry = st.text_area(
            "Morning pages — empty your head onto the page",
            value=e.get("morning_entry") or "", height=220, key="m_entry",
            placeholder=("Good morning. Write freely — how you slept, what's on your mind, "
                         "what you're worried or excited about. No structure, no judging. "
                         "This is the raw feed of your life."))
        _mw = len(morning_entry.split())
        st.caption(f"✍️ {_mw} words")

        if st.button("💾 Save Morning", type="primary", key="save_morning"):
            ok, err = _save_partial({
                "waking_mood": waking_mood,
                "sleep_duration": sleep_duration,
                "dream_log": dream_log.strip(),
                "intention": intention.strip(),
                "morning_entry": morning_entry.strip(),
                "morning_saved_at": "now()",
            })
            if ok:
                st.toast("🌅 Morning saved. Go win the day.")
                st.rerun()
            else:
                st.error(f"Save failed: {err}")

    # ============================================================
    # 🌙 EVENING — the 10–15 min self-talk
    # ============================================================
    with tab_evening:
        if e.get("intention"):
            st.markdown(f"""
            <div style="border-left:3px solid #facc15;padding:0.5rem 1rem;font-size:0.9rem;color:#e8e0ff;">
                🎯 This morning you said: <b>“{e['intention']}”</b> — did it happen?
            </div><br>""", unsafe_allow_html=True)

        # ---- 15-minute self-talk timer ----
        with st.expander("⏱ Start your 15-minute reflection timer"):
            components.html("""
            <div style="font-family:monospace;text-align:center;background:transparent;">
              <div id="t" style="font-size:2.6rem;font-weight:700;color:#a78bfa;">15:00</div>
              <button onclick="startT()" style="margin:4px;padding:6px 22px;background:#1a1a3a;color:#e8e0ff;border:1px solid #a78bfa;border-radius:8px;cursor:pointer;">Start</button>
              <button onclick="resetT()" style="margin:4px;padding:6px 22px;background:#1a1a3a;color:#9ca3af;border:1px solid #444;border-radius:8px;cursor:pointer;">Reset</button>
              <script>
                let secs=900, iv=null;
                function draw(){const m=String(Math.floor(secs/60)).padStart(2,'0'),s=String(secs%60).padStart(2,'0');
                  const el=document.getElementById('t'); el.textContent=m+':'+s;
                  if(secs===0){el.textContent='✦ done — save your entry'; el.style.color='#4ade80';}}
                function startT(){if(iv)return; iv=setInterval(()=>{if(secs>0){secs--;draw();}else{clearInterval(iv);iv=null;draw();}},1000);}
                function resetT(){clearInterval(iv);iv=null;secs=900;
                  document.getElementById('t').style.color='#a78bfa';draw();}
              </script>
            </div>""", height=130)

        st.markdown(f"""
        <div class="card" style="border-color:#a78bfa;">
            <div style="font-size:0.7rem;color:#3a3a6a;font-family:'JetBrains Mono',monospace;letter-spacing:1px;">TONIGHT'S PROMPT</div>
            <div class="chivalry" style="font-size:1.05rem;margin-top:0.4rem;"><em>{_prompt_today}</em></div>
        </div>""", unsafe_allow_html=True)

        evening_entry = st.text_area(
            "Evening reflection — talk to yourself honestly",
            value=e.get("evening_entry") or "", height=300, key="e_entry",
            placeholder=("Sit with the day for 10–15 minutes. What actually happened? "
                         "What did you feel? Where were you strong, where did you fold? "
                         "End by telling yourself something true and encouraging — "
                         "you're the one reading this in a year."))
        _ew = len(evening_entry.split())
        st.caption(f"✍️ {_ew} words" + (" · that's a real entry 💪" if _ew >= 150 else ""))

        gratitude = st.text_input("🙏 One thing you're grateful for today",
                                  value=e.get("gratitude") or "", key="e_gratitude",
                                  placeholder="Small counts. Plov counts.")

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.75rem;color:#555;font-family:\'JetBrains Mono\',monospace;margin-bottom:0.8rem;letter-spacing:1px;">STATE CHECK</div>', unsafe_allow_html=True)

        col1, col2 = st.columns([2, 1])
        with col1:
            mood_score = st.slider("Mood Score", 1, 10, int(e.get("mood_score") or 5), key="e_mood")
        with col2:
            st.markdown(f"<br>{mood_badge(mood_score)}", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            clarity_options = ["foggy", "normal", "sharp"]
            _cd = clarity_options.index(e.get("mental_clarity")) if e.get("mental_clarity") in clarity_options else 1
            mental_clarity = st.selectbox("Mental Clarity", clarity_options, index=_cd, key="e_clarity")
        with col2:
            dominant_emotion = st.text_input("Dominant Emotion", value=e.get("dominant_emotion") or "",
                                             placeholder="e.g. anxious, hopeful, calm…", key="e_emotion")

        col1, col2 = st.columns(2)
        with col1:
            phys_options = ["tired", "neutral", "energized"]
            _pd = phys_options.index(e.get("physical_state")) if e.get("physical_state") in phys_options else 1
            physical_state = st.selectbox("Physical State", phys_options, index=_pd, key="e_phys")
        with col2:
            ment_options = ["calm", "stable", "stressed", "heavy"]
            _md = ment_options.index(e.get("mental_state")) if e.get("mental_state") in ment_options else 1
            mental_state = st.selectbox("Mental State", ment_options, index=_md, key="e_ment")

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.75rem;color:#555;font-family:\'JetBrains Mono\',monospace;margin-bottom:0.8rem;letter-spacing:1px;">THE DAY ITSELF</div>', unsafe_allow_html=True)

        daily_summary = st.text_area("Daily Summary (facts — what happened)",
                                     value=e.get("daily_summary") or "", height=90, key="e_summary")
        good_deed = st.text_input("Good Deed of the Day", value=e.get("good_deed") or "",
                                  placeholder="One good thing you did for someone…", key="e_deed")

        st.markdown("**Accomplishments**")
        acc_col1, acc_col2 = st.columns([4, 1])
        with acc_col1:
            new_acc = st.text_input("Add accomplishment", key="new_acc_input",
                                    label_visibility="collapsed", placeholder="What did you get done?")
        with acc_col2:
            if st.button("＋ Add", key="add_acc"):
                if new_acc.strip():
                    st.session_state.accomplishments.append(new_acc.strip())
                    st.rerun()
        for i, acc in enumerate(st.session_state.accomplishments):
            c1, c2 = st.columns([6, 1])
            with c1:
                st.markdown(f"&nbsp; ✦ {acc}")
            with c2:
                if st.button("✕", key=f"del_acc_{i}"):
                    st.session_state.accomplishments.pop(i)
                    st.rerun()

        if st.button("💾 Save Evening", type="primary", key="save_evening"):
            fields = {
                "evening_entry": evening_entry.strip(),
                "gratitude": gratitude.strip(),
                "mood_score": mood_score,
                "mental_clarity": mental_clarity,
                "dominant_emotion": dominant_emotion.strip(),
                "physical_state": physical_state,
                "mental_state": mental_state,
                "daily_summary": daily_summary.strip(),
                "good_deed": good_deed.strip(),
                "accomplishments": st.session_state.accomplishments,
                "evening_saved_at": "now()",
            }
            # keep old column populated so the Archive page keeps working
            fields["self_assessment"] = evening_entry.strip()[:1000]
            ok, err = _save_partial(fields)
            if ok:
                st.toast("🌙 Evening saved. Rest well — the archive remembers.")
                st.rerun()
            else:
                st.error(f"Save failed: {err}")

    # ============================================================
    # 📊 METRICS — spending, media, training
    # ============================================================
    with tab_metrics:
        col1, col2 = st.columns(2)
        with col1:
            phone_off = st.time_input("Phone Off Time", value=None, key="mt_phone_off")
            daily_spending = st.number_input("Daily Spending Total (UZS)", min_value=0.0, step=1000.0,
                                             value=float(e.get("daily_spending") or 0.0), key="mt_spend")
        with col2:
            phone_on = st.time_input("Phone On Time", value=None, key="mt_phone_on")
            spending_notes = st.text_input("Spending Notes", value=e.get("spending_notes") or "",
                                           placeholder="What did you spend on?", key="mt_spend_notes")

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.75rem;color:#555;font-family:\'JetBrains Mono\',monospace;margin-bottom:0.8rem;letter-spacing:1px;">MEDIA CONSUMED</div>', unsafe_allow_html=True)
        med_col1, med_col2 = st.columns([4, 1])
        with med_col1:
            new_media = st.text_input("Add media", key="new_media_input", label_visibility="collapsed",
                                      placeholder="Anime / movie / series title…")
        with med_col2:
            if st.button("＋ Add", key="add_media"):
                if new_media.strip():
                    st.session_state.media_list.append(new_media.strip())
                    st.rerun()
        for i, m in enumerate(st.session_state.media_list):
            c1, c2 = st.columns([6, 1])
            with c1:
                st.markdown(f"&nbsp; 🎬 {m}")
            with c2:
                if st.button("✕", key=f"del_media_{i}"):
                    st.session_state.media_list.pop(i)
                    st.rerun()

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.75rem;color:#555;font-family:\'JetBrains Mono\',monospace;margin-bottom:0.8rem;letter-spacing:1px;">EXERCISE & TRAINING</div>', unsafe_allow_html=True)

        rest_day = st.checkbox("🛌 Rest day — no training today", value=bool(e.get("rest_day")), key="mt_rest")
        if not rest_day:
            WORKOUT_TYPES = ["🥊 Boxing / MMA", "🏃 Cardio", "🏋️ Weights", "🧘 Mobility / Stretching",
                             "⚽ Sport", "🤸 Calisthenics", "🥋 Martial Arts", "Other"]
            st.markdown("**Sessions**")
            w_col1, w_col2, w_col3 = st.columns([2, 3, 1])
            with w_col1:
                new_w_type = st.selectbox("Type", WORKOUT_TYPES, key="new_w_type", label_visibility="collapsed")
            with w_col2:
                new_w_label = st.text_input("Session label", key="new_w_label", label_visibility="collapsed",
                                            placeholder="e.g. Sparring 3×3, Bench 4×8…")
            with w_col3:
                if st.button("＋ Add", key="add_workout"):
                    if new_w_label.strip():
                        st.session_state.workouts.append({"type": new_w_type, "label": new_w_label.strip()})
                        st.rerun()
            for i, w in enumerate(st.session_state.workouts):
                wc1, wc2 = st.columns([6, 1])
                with wc1:
                    st.markdown(f"&nbsp; {w['type']} &nbsp; `{w['label']}`")
                with wc2:
                    if st.button("✕", key=f"del_workout_{i}"):
                        st.session_state.workouts.pop(i)
                        st.rerun()

            col1, col2, col3 = st.columns(3)
            with col1:
                training_duration = st.number_input("Total duration (min)", 0, 480, int(e.get("training_duration") or 0),
                                                    step=5, key="mt_dur")
            with col2:
                intensity_opts = ["easy", "moderate", "hard", "max"]
                _id = intensity_opts.index(e.get("training_intensity")) if e.get("training_intensity") in intensity_opts else 1
                training_intensity = st.selectbox("Intensity", intensity_opts, index=_id, key="mt_int")
            with col3:
                body_feel_opts = ["great", "good", "okay", "sore", "injured"]
                _bd = body_feel_opts.index(e.get("body_feel")) if e.get("body_feel") in body_feel_opts else 1
                body_feel = st.selectbox("Body Feel After", body_feel_opts, index=_bd, key="mt_body")
            training_notes = st.text_input("Training Notes", value=e.get("training_notes") or "",
                                           placeholder="PRs, observations, what to improve…", key="mt_notes")
        else:
            training_duration, training_intensity, body_feel, training_notes = 0, None, None, ""

        # training streak
        try:
            _streak_logs = supabase.table("daily_logs").select("date,rest_day,training_duration") \
                .order("date", desc=True).limit(60).execute().data or []
            _tstreak = 0
            for _sl in _streak_logs:
                if not _sl.get("rest_day") and (_sl.get("training_duration") or 0) > 0:
                    _tstreak += 1
                else:
                    break
            if _tstreak >= 2:
                st.markdown(f'<div style="font-size:0.95rem;color:#facc15;font-family:JetBrains Mono,monospace;">🔥 {_tstreak} day training streak</div>', unsafe_allow_html=True)
        except Exception:
            pass

        if st.button("💾 Save Metrics", type="primary", key="save_metrics"):
            ok, err = _save_partial({
                "phone_off_time": str(phone_off) if phone_off else None,
                "phone_on_time": str(phone_on) if phone_on else None,
                "daily_spending": daily_spending,
                "spending_notes": spending_notes.strip(),
                "media_consumed": st.session_state.media_list,
                "rest_day": rest_day,
                "workouts": st.session_state.workouts if not rest_day else [],
                "training_duration": training_duration if not rest_day else 0,
                "training_intensity": training_intensity if not rest_day else None,
                "body_feel": body_feel if not rest_day else None,
                "training_notes": training_notes.strip() if not rest_day else "",
            })
            if ok:
                st.toast("📊 Metrics saved.")
                st.rerun()
            else:
                st.error(f"Save failed: {err}")

    # ============================================================
    # 📤 Share today's entry (only once something is saved)
    # ============================================================
    if e:
        with st.expander("📤 Share / export today"):
            def build_share_text(entry):
                d = entry.get("date", str(today))
                lines = [f"📅 *Life Archive — {d}*", ""]
                if entry.get("intention"):
                    lines += [f"🎯 Intention: {entry['intention']}"]
                lines += [f"🧠 Mood: {entry.get('mood_score','—')}/10  |  Clarity: {entry.get('mental_clarity','—')}  |  Emotion: {entry.get('dominant_emotion','—')}",
                          f"😴 Sleep: {entry.get('sleep_duration','—')}h", ""]
                if entry.get("daily_summary"):
                    lines += ["📝 *Summary*", entry["daily_summary"], ""]
                accs = entry.get("accomplishments") or []
                if accs:
                    lines += ["✦ *Accomplishments*"] + [f"  • {a}" for a in accs] + [""]
                if entry.get("gratitude"):
                    lines += [f"🙏 Grateful for: {entry['gratitude']}", ""]
                if entry.get("good_deed"):
                    lines += [f"💛 Good deed: {entry['good_deed']}", ""]
                if entry.get("rest_day"):
                    lines += ["🛌 Rest day", ""]
                elif entry.get("training_duration"):
                    lines += [f"🏋️ Training — {entry.get('training_duration',0)} min | {(entry.get('training_intensity') or '').upper()}"]
                    for w in (entry.get("workouts") or []):
                        lines.append(f"  {w.get('type','')} {w.get('label','')}")
                    lines.append("")
                media = entry.get("media_consumed") or []
                if media:
                    lines += ["🎬 *Media*"] + [f"  • {m}" for m in media] + [""]
                lines.append("— sent from Life Archive")
                return "\n".join(lines)

            share_text = build_share_text(e)
            st.text_area("Copy to Telegram or your channel", value=share_text, height=280, key="share_export_box")
            _dl_col, _tg_col = st.columns([1, 2])
            with _dl_col:
                import json as _json
                st.download_button("⬇️ Download as JSON", data=_json.dumps(e, indent=2, default=str),
                                   file_name=f"life_archive_{e.get('date', str(today))}.json",
                                   mime="application/json", use_container_width=True)
            with _tg_col:
                tg_url = f"https://t.me/share/url?url=&text={share_text.replace(' ','%20').replace(chr(10),'%0A')[:2000]}"
                st.markdown(f'<a href="{tg_url}" target="_blank" style="display:inline-block;width:100%;margin-top:0.3rem;padding:0.4rem 1.2rem;background:#2a2a4a;border:1px solid #a78bfa;border-radius:8px;color:#e8e0ff;text-decoration:none;font-family:JetBrains Mono,monospace;font-size:0.8rem;text-align:center;">📨 Open in Telegram</a>', unsafe_allow_html=True)

elif page == "📖  Reading Log":
    st.markdown('<div class="section-header">📖 Reading Log</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">track every book · session by session</div>', unsafe_allow_html=True)

    # ---- Fetch all books ----
    try:
        all_books = supabase.table("books").select("*").order("created_at", desc=False).execute().data or []
        reading_books   = [b for b in all_books if b["status"] == "reading"]
        finished_books  = [b for b in all_books if b["status"] == "finished"]
    except Exception:
        all_books = reading_books = finished_books = []

    # ---- Tabs ----
    tab1, tab2, tab3 = st.tabs(["📚 Currently Reading", "✅ Finished Shelf", "➕ Add New Book"])

    # ============================================================
    # TAB 1 — CURRENTLY READING
    # ============================================================
    with tab1:
        if not reading_books:
            st.info("No books in progress. Add one in the 'Add New Book' tab.")
        else:
            for book in reading_books:
                # Fetch sessions for this book
                try:
                    sessions = supabase.table("reading_sessions").select("*").eq("book_id", book["id"]).order("session_date", desc=False).execute().data or []
                except Exception:
                    sessions = []

                total_read = sum(s["pages_read"] for s in sessions)
                total_pages = book["total_pages"]
                progress = min(int((total_read / total_pages) * 100), 100) if total_pages > 0 else 0

                # Inactivity check
                inactive_warning = ""
                if sessions:
                    from datetime import datetime, timedelta
                    last_session_date = datetime.strptime(sessions[-1]["session_date"], "%Y-%m-%d").date()
                    days_since = (date.today() - last_session_date).days
                    if days_since >= 7:
                        inactive_warning = f'<span class="badge badge-red">🔴 {days_since} days inactive</span>'

                # Progress circle via SVG
                radius = 36
                circumference = 2 * 3.14159 * radius
                stroke_offset = circumference * (1 - progress / 100)
                circle_color = "#4ade80" if progress >= 75 else "#facc15" if progress >= 40 else "#60a5fa"

                col1, col2 = st.columns([1, 3])
                with col1:
                    st.markdown(f"""
                    <div style="display:flex;justify-content:center;align-items:center;padding:1rem 0;">
                        <svg width="100" height="100" viewBox="0 0 100 100">
                            <circle cx="50" cy="50" r="{radius}" fill="none" stroke="#2a2a2a" stroke-width="8"/>
                            <circle cx="50" cy="50" r="{radius}" fill="none" stroke="{circle_color}" stroke-width="8"
                                stroke-dasharray="{circumference:.1f}"
                                stroke-dashoffset="{stroke_offset:.1f}"
                                stroke-linecap="round"
                                transform="rotate(-90 50 50)"/>
                            <text x="50" y="54" text-anchor="middle" fill="#fff"
                                font-family="JetBrains Mono" font-size="14" font-weight="700">{progress}%</text>
                        </svg>
                    </div>
                    """, unsafe_allow_html=True)

                with col2:
                    st.markdown(f"### {book['title']}")
                    if inactive_warning:
                        st.markdown(inactive_warning, unsafe_allow_html=True)
                    genre_tag = f'<span class="badge badge-grey">{book["genre"]}</span> &nbsp;' if book.get("genre") else ""
                    # Estimated finish date
                    est_finish = ""
                    if sessions and total_pages > total_read:
                        avg_pages = sum(s["pages_read"] for s in sessions) / len(sessions)
                        if avg_pages > 0:
                            days_left = int((total_pages - total_read) / avg_pages)
                            from datetime import timedelta
                            finish_date = (date.today() + timedelta(days=days_left)).strftime("%b %d")
                            est_finish = f'<span class="badge badge-blue">📅 est. finish {finish_date}</span>'
                    st.markdown(f"{genre_tag}{est_finish}", unsafe_allow_html=True)
                    st.caption(f"{book.get('author','Unknown')} · {total_read} / {total_pages} pages · {len(sessions)} sessions")

                    if sessions:
                        with st.expander(f"🗂 {len(sessions)} past sessions", expanded=st.session_state.get(f"show_past_{book['id']}", False)):
                            for s in reversed(sessions):
                                stars = '★' * s['session_rating'] + '☆' * (5 - s['session_rating'])
                                st.markdown(f"""
                                <div class="card" style="margin-bottom:0.4rem;">
                                    <div style="display:flex;justify-content:space-between;">
                                        <span style="font-family:'JetBrains Mono',monospace;font-size:0.8rem;color:#888;">{s['session_date']}</span>
                                        <span class="badge badge-blue">+{s['pages_read']} pages</span>
                                        <span class="badge badge-yellow">{stars}</span>
                                    </div>
                                    <div style="margin-top:0.4rem;color:#ccc;font-size:0.85rem;">{s.get('learned','—')}</div>
                                    {f'<div style="margin-top:0.4rem;padding:0.3rem 0.6rem;border-left:3px solid #a78bfa;color:#a78bfa;font-size:0.82rem;font-style:italic;">{s["highlight"]}</div>' if s.get("highlight") else ""}
                                </div>
                                """, unsafe_allow_html=True)

                    with st.expander(f"➕ Log new session — '{book['title']}'", expanded=not st.session_state.get(f"show_past_{book['id']}", False)):
                        s_pages   = st.number_input("Pages read today", min_value=1, step=1, key=f"pages_{book['id']}")
                        s_learned = st.text_area("What did you learn?", key=f"learned_{book['id']}", height=80)
                        s_highlight = st.text_input("Quote or highlight (optional)", key=f"highlight_{book['id']}", placeholder="A passage worth saving...")
                        s_rating  = st.slider("How was the session?", 1, 5, 3, key=f"srating_{book['id']}")

                        if st.button("Save Session", key=f"save_session_{book['id']}"):
                            try:
                                supabase.table("reading_sessions").insert({
                                    "book_id": book["id"],
                                    "session_date": str(date.today()),
                                    "pages_read": s_pages,
                                    "learned": s_learned.strip(),
                                    "highlight": s_highlight.strip(),
                                    "session_rating": s_rating
                                }).execute()
                                new_total = total_read + s_pages
                                if new_total >= total_pages:
                                    st.session_state[f"complete_{book['id']}"] = True
                                    st.success("🎉 You finished the book!")
                                else:
                                    st.success(f"✅ Session saved! {total_pages - new_total} pages left.")
                                st.session_state[f"show_past_{book['id']}"] = True
                                st.rerun()
                            except Exception as ex:
                                st.error(f"Error: {ex}")

                    # Completion form (auto-triggered)
                    if st.session_state.get(f"complete_{book['id']}", False):
                        st.markdown("---")
                        st.markdown("### 🏁 Book Complete — Final Review")
                        f_rating  = st.slider("Overall rating (1–10)", 1, 10, 8, key=f"frating_{book['id']}")
                        f_review  = st.text_area("How was the book overall?", key=f"freview_{book['id']}", height=80)
                        f_rec     = st.radio("Would you recommend it?", ["Yes", "No"], key=f"frec_{book['id']}")
                        f_takeaway = st.text_input("Key takeaway", key=f"ftakeaway_{book['id']}")

                        if st.button("✅ Mark as Finished", key=f"finish_{book['id']}"):
                            try:
                                supabase.table("books").update({
                                    "status": "finished",
                                    "overall_rating": f_rating,
                                    "review": f_review.strip(),
                                    "recommend": f_rec == "Yes",
                                    "key_takeaway": f_takeaway.strip(),
                                    "date_finished": str(date.today())
                                }).eq("id", book["id"]).execute()
                                st.session_state[f"complete_{book['id']}"] = False
                                st.success("Book moved to finished shelf!")
                                st.rerun()
                            except Exception as ex:
                                st.error(f"Error: {ex}")

                st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # TAB 2 — FINISHED SHELF
    # ============================================================
    with tab2:
        if not finished_books:
            st.info("No finished books yet. Keep reading!")
        else:
            _min_rating = st.selectbox("Minimum rating", ["All", "≥ 5", "≥ 7", "≥ 9"],
                                       label_visibility="collapsed", key="shelf_rating_filter")
            _rating_floor = {"All": 0, "≥ 5": 5, "≥ 7": 7, "≥ 9": 9}[_min_rating]
            _shown = [b for b in finished_books
                      if (_rating_floor == 0 or (isinstance(b.get("overall_rating"), (int,float)) and b["overall_rating"] >= _rating_floor))]
            if not _shown:
                st.info(f"No finished books with rating {_min_rating}.")
            for book in _shown:
                rec_badge = '<span class="badge badge-green">👍 Recommended</span>' if book.get("recommend") else '<span class="badge badge-grey">👎 Not Recommended</span>'
                rating = book.get("overall_rating")
                rating = rating if isinstance(rating, (int, float)) else "—"
                rating_color = "#4ade80" if (rating != "—" and rating >= 8) else "#facc15" if (rating != "—" and rating >= 5) else "#f87171"
                genre_badge = f'<span class="badge badge-grey">{book["genre"]}</span>' if book.get("genre") else ""

                st.markdown(f"""
                <div class="card">
                    <div style="display:flex;justify-content:space-between;align-items:center;">
                        <div>
                            <span style="font-size:1rem;font-weight:800;color:#fff;">{book['title']}</span>
                            <span style="font-size:0.8rem;color:#666;margin-left:0.8rem;font-family:'JetBrains Mono',monospace;">{book.get('author','')}</span>
                        </div>
                        <div style="display:flex;gap:0.5rem;align-items:center;">
                            {genre_badge}
                            <span style="font-family:'JetBrains Mono',monospace;font-size:1.2rem;font-weight:800;color:{rating_color};">{rating}/10</span>
                            {rec_badge}
                        </div>
                    </div>
                    <div style="margin-top:0.6rem;color:#aaa;font-size:0.85rem;">{book.get('review','')}</div>
                    <div style="margin-top:0.4rem;font-size:0.8rem;color:#555;font-family:'JetBrains Mono',monospace;">
                        💡 {book.get('key_takeaway','—')} &nbsp;·&nbsp; finished {book.get('date_finished','—')}
                    </div>
                </div>
                """, unsafe_allow_html=True)

    # ============================================================
    # TAB 3 — ADD NEW BOOK
    # ============================================================
    with tab3:
        st.markdown("#### Add a New Book")
        b_title  = st.text_input("Book Title", placeholder="e.g. Limitless")
        b_author = st.text_input("Author", placeholder="e.g. Jim Kwik")
        b_pages  = st.number_input("Total Pages", min_value=1, step=1, value=300)
        b_genres = ["Fiction", "Non-Fiction", "Fantasy", "Sci-Fi", "Mystery", "History",
                    "Biography", "Self-Help", "Economics", "Philosophy", "Science", "Other"]
        b_genre  = st.selectbox("Genre", b_genres)

        if st.button("➕ Add Book", type="primary"):
            if not b_title.strip():
                st.error("Please enter a title.")
            else:
                try:
                    supabase.table("books").insert({
                        "title": b_title.strip(),
                        "author": b_author.strip(),
                        "total_pages": b_pages,
                        "genre": b_genre,
                        "status": "reading",
                        "date_started": str(date.today())
                    }).execute()
                    st.success(f"'{b_title}' added to your reading list!")
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

elif page == "🚨  Life Event":
    st.markdown('<div class="section-header">🚨 Life Event Diary</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">record moments that actually mattered</div>', unsafe_allow_html=True)

    if "life_event_view" not in st.session_state:
        st.session_state["life_event_view"] = "➕ Log New Event"

    le_options = ["📋 Past Events", "➕ Log New Event"]
    le_index = le_options.index(st.session_state["life_event_view"]) if st.session_state["life_event_view"] in le_options else 0
    view = st.radio("Life Event View", le_options, index=le_index,
                    horizontal=True, label_visibility="collapsed",
                    key=f"life_event_radio_{le_index}")
    st.session_state["life_event_view"] = view
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # PAST EVENTS
    # ============================================================
    if view == "📋 Past Events":
        event_types = ["All", "movie", "accident", "realization", "personal event", "other"]
        filter_type = st.selectbox("Filter by type", event_types, label_visibility="collapsed")

        try:
            query = supabase.table("life_events").select("*").order("event_date", desc=True)
            if filter_type != "All":
                query = query.eq("event_type", filter_type)
            events = query.execute().data or []
        except Exception:
            events = []

        if not events:
            st.info("No life events logged yet.")
        else:
            for e in events:
                sig = e.get("significance_score", 1)
                if sig >= 5:
                    border_color = "#fbbf24"
                elif sig >= 4:
                    border_color = "#fb923c"
                elif sig >= 3:
                    border_color = "#facc15"
                else:
                    border_color = "#444"

                type_colors = {
                    "realization":    "badge-green",
                    "milestone":      "badge-green",
                    "personal event": "badge-blue",
                    "conversation":   "badge-blue",
                    "decision":       "badge-yellow",
                    "travel":         "badge-yellow",
                    "accident":       "badge-red",
                    "other":          "badge-grey"
                }
                type_badge = f'<span class="badge {type_colors.get(e["event_type"], "badge-grey")}">{e["event_type"]}</span>'
                people_html = f'<div style="margin-top:0.3rem;color:#888;font-size:0.82rem;">👥 {e["people_involved"]}</div>' if e.get('people_involved') else ''
                change_html = f'<div style="margin-top:0.3rem;color:#666;font-size:0.82rem;font-style:italic;">🔄 {e["would_change"]}</div>' if e.get('would_change') else ''

                le_card = (
                    f'<div class="card" style="border-left: 3px solid {border_color}; margin-bottom:1rem;">'
                    '<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">'
                    f'<span style="font-size:1rem;font-weight:800;color:#fff;">{e["event_title"]}</span>'
                    f'<div style="display:flex;gap:0.5rem;align-items:center;">{type_badge} {significance_badge(sig)}</div>'
                    '</div>'
                    f'<div style="color:#ccc;font-size:0.9rem;margin-bottom:0.5rem;">{e.get("event_description","")}</div>'
                    f'<div style="color:#888;font-size:0.85rem;font-style:italic;">💭 {e.get("emotional_impact","")}</div>'
                    f'{people_html}{change_html}'
                    f'<div style="margin-top:0.5rem;font-size:0.75rem;color:#555;font-family:\'JetBrains Mono\',monospace;">{e.get("event_date","")}</div>'
                    '</div>'
                )
                st.markdown(le_card, unsafe_allow_html=True)
                col_edit, col_del = st.columns([1, 1])
                with col_edit:
                    if st.button("✏️ Edit", key=f"edit_event_{e['id']}"):
                        st.session_state["editing_event"] = e
                        st.session_state["life_event_view"] = "➕ Log New Event"
                        st.rerun()
                with col_del:
                    if st.button("🗑️ Delete", key=f"del_event_{e['id']}"):
                        st.session_state[f"confirm_del_{e['id']}"] = True
                        st.rerun()
                if st.session_state.get(f"confirm_del_{e['id']}", False):
                    st.warning(f"Delete **{e['event_title']}**? This cannot be undone.")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Yes, delete", key=f"yes_del_{e['id']}"):
                            supabase.table("life_events").delete().eq("id", e["id"]).execute()
                            st.session_state[f"confirm_del_{e['id']}"] = False
                            st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_del_{e['id']}"):
                            st.session_state[f"confirm_del_{e['id']}"] = False
                            st.rerun()

    # ============================================================
    # LOG NEW EVENT
    # ============================================================
    else:
        editing = st.session_state.get("editing_event", None)
        st.markdown(f"#### {'✏️ Editing Event' if editing else 'What happened?'}")

        e_types_list = ["personal event", "milestone", "realization", "decision",
                        "conversation", "travel", "accident", "other"]
        e_title  = st.text_input("Event Title", value=editing["event_title"] if editing else "", placeholder="e.g. Got accepted...")
        e_type   = st.selectbox("Event Type", e_types_list, index=e_types_list.index(editing["event_type"]) if editing and editing.get("event_type") in e_types_list else 0)
        e_desc   = st.text_area("Description", value=editing.get("event_description","") if editing else "", height=100)
        e_impact = st.text_area("Emotional Impact", value=editing.get("emotional_impact","") if editing else "", height=80)
        e_people = st.text_input("People Involved (optional)", value=editing.get("people_involved","") if editing else "",
                                 placeholder="e.g. Dad, Professor Kim, my team...")
        e_change = st.text_input("What would you do differently? (optional)", value=editing.get("would_change","") if editing else "",
                                 placeholder="Hindsight reflection...")

        col1, col2 = st.columns(2)
        with col1:
            e_sig = st.slider("Significance (1–5)", 1, 5, min(5, max(1, editing["significance_score"])) if editing and editing.get("significance_score") else 3)
            st.markdown(significance_badge(e_sig), unsafe_allow_html=True)
        with col2:
            from datetime import datetime
            default_date = datetime.strptime(editing["event_date"], "%Y-%m-%d").date() if editing else date.today()
            e_date = st.date_input("Date", value=default_date)

        if st.button("💾 Update Event" if editing else "💾 Save Event", type="primary"):
            if not e_title.strip():
                st.error("Please enter an event title.")
            else:
                try:
                    record = {
                        "event_title": e_title.strip(),
                        "event_type": e_type,
                        "event_description": e_desc.strip(),
                        "emotional_impact": e_impact.strip(),
                        "people_involved": e_people.strip(),
                        "would_change": e_change.strip(),
                        "significance_score": e_sig,
                        "event_date": str(e_date)
                    }
                    if editing:
                        supabase.table("life_events").update(record).eq("id", editing["id"]).execute()
                        st.success("✅ Event updated.")
                        st.session_state["editing_event"] = None
                    else:
                        supabase.table("life_events").insert(record).execute()
                        st.success("✅ Event saved.")
                    st.session_state["life_event_view"] = "📋 Past Events"
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

        if editing and st.button("Cancel"):
            st.session_state["editing_event"] = None
            st.rerun()
                
elif page == "🛍️  Purchase Tracker":
    st.markdown('<div class="section-header">🛍️ Purchase Tracker</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">track what you buy · review what it was worth</div>', unsafe_allow_html=True)

    if "purchase_view" not in st.session_state:
        st.session_state["purchase_view"] = "➕ Add Purchase"
    if "editing_purchase" not in st.session_state:
        st.session_state["editing_purchase"] = None

    p_options = ["🔴 Pending Review", "📋 All Purchases", "➕ Add Purchase"]
    p_index = p_options.index(st.session_state["purchase_view"]) if st.session_state["purchase_view"] in p_options else 0
    view = st.radio("Purchase View", p_options, index=p_index,
                    horizontal=True, label_visibility="collapsed",
                    key=f"purchase_radio_{p_index}")
    st.session_state["purchase_view"] = view
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # PENDING REVIEW
    # ============================================================
    if view == "🔴 Pending Review":
        try:
            all_purchases = supabase.table("purchases").select("*").eq("phase2_completed", False).order("purchase_date", desc=True).execute().data or []
            pending = [p for p in all_purchases if p.get("review_due_date") and p["review_due_date"] <= str(date.today())]
        except Exception:
            pending = []

        if not pending:
            st.success("✅ No purchases pending review.")
        else:
            st.markdown(f'<div class="section-sub">🔴 {len(pending)} purchase(s) waiting for your honest review</div>', unsafe_allow_html=True)
            for p in pending:
                days_overdue = (date.today() - date.fromisoformat(p["review_due_date"])).days
                currency = p.get("currency", "UZS")
                st.markdown(f"""
                <div class="card" style="border-left: 3px solid #f87171;">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.4rem;">
                        <span style="font-weight:800;color:#fff;font-size:1rem;">{p['item_name']}</span>
                        <div style="display:flex;gap:0.5rem;">
                            <span class="badge badge-red">🔴 {days_overdue}d overdue</span>
                            <span class="badge badge-grey">{p.get('category','')}</span>
                        </div>
                    </div>
                    <div style="font-size:0.85rem;color:#888;font-family:'JetBrains Mono',monospace;">
                        {p.get('amount',0):,.0f} {currency} · bought {p.get('purchase_date','')} · expected: {p.get('expected_value','')}
                    </div>
                </div>
                """, unsafe_allow_html=True)

                with st.expander(f"📝 Review '{p['item_name']}'"):
                    r_usage      = st.selectbox("How often did you use it?", ["never", "once", "few times", "regularly", "daily"], key=f"usage_{p['id']}")
                    r_actual     = st.text_input("Actual use case", key=f"actual_{p['id']}", placeholder="What did you actually use it for?")
                    r_sat        = st.slider("Satisfaction (1–5)", 1, 5, 3, key=f"sat_{p['id']}")
                    r_regret     = st.slider("Regret level (1–5)", 1, 5, 2, key=f"regret_{p['id']}")
                    r_worth      = st.radio("Worth it?", ["Yes", "No"], key=f"worth_{p['id']}", horizontal=True)
                    r_reflection = st.text_area("Reflection", key=f"reflection_{p['id']}", height=80, placeholder="What did you learn from this purchase?")

                    if st.button("✅ Submit Review", key=f"review_{p['id']}", type="primary"):
                        try:
                            supabase.table("purchases").update({
                                "usage_frequency": r_usage,
                                "actual_use_case": r_actual.strip(),
                                "satisfaction_level": r_sat,
                                "regret_level": r_regret,
                                "worth_it": r_worth == "Yes",
                                "review_reflection": r_reflection.strip(),
                                "phase2_completed": True
                            }).eq("id", p["id"]).execute()
                            st.success("✅ Review saved.")
                            st.rerun()
                        except Exception as ex:
                            st.error(f"Error: {ex}")

    # ============================================================
    # ALL PURCHASES
    # ============================================================
    elif view == "📋 All Purchases":
        try:
            all_purchases = supabase.table("purchases").select("*").order("purchase_date", desc=True).execute().data or []
        except Exception:
            all_purchases = []

        if not all_purchases:
            st.info("No purchases logged yet.")
        else:
            # Per-currency monthly summary
            from datetime import timedelta as _td
            _this_month = date.today().strftime("%Y-%m")
            _last_month = (date.today().replace(day=1) - _td(days=1)).strftime("%Y-%m")

            _currencies_seen = sorted(set(
                p.get("currency","UZS") for p in all_purchases
                if str(p.get("purchase_date","")).startswith(_this_month)
                or str(p.get("purchase_date","")).startswith(_last_month)
            ))

            if _currencies_seen:
                _cur_cols = st.columns(len(_currencies_seen))
                for _ci, _cur in enumerate(_currencies_seen):
                    _this_t = sum(p.get("amount",0) or 0 for p in all_purchases
                                  if p.get("currency")==_cur and str(p.get("purchase_date","")).startswith(_this_month))
                    _last_t = sum(p.get("amount",0) or 0 for p in all_purchases
                                  if p.get("currency")==_cur and str(p.get("purchase_date","")).startswith(_last_month))
                    _diff   = _this_t - _last_t
                    _diff_color = "#f87171" if _diff > 0 else "#4ade80"
                    _diff_str   = f'{"+" if _diff>0 else ""}{_diff:,.0f}' if _last_t else "new"
                    with _cur_cols[_ci]:
                        st.markdown(f"""
                        <div class="card" style="text-align:center;">
                            <div style="font-size:0.7rem;color:#555;font-family:'JetBrains Mono',monospace;letter-spacing:1px;">THIS MONTH · {_cur}</div>
                            <div style="font-size:1.5rem;font-weight:800;color:#e8e0ff;font-family:'JetBrains Mono',monospace;">{_this_t:,.0f}</div>
                            <div style="font-size:0.8rem;color:{_diff_color};font-family:'JetBrains Mono',monospace;">{_diff_str} vs last month</div>
                        </div>
                        """, unsafe_allow_html=True)

            # This-month category breakdown (by number of purchases — currency-agnostic)
            _month_p = [p for p in all_purchases if str(p.get("purchase_date","")).startswith(_this_month)]
            if _month_p:
                from collections import Counter as _Counter
                _cat_counts = _Counter(p.get("category","other") for p in _month_p)
                _max_c = max(_cat_counts.values())
                _cat_colors = {"food":"#4ade80","clothing":"#c084fc","tech":"#818cf8",
                               "entertainment":"#facc15","transport":"#60a5fa","other":"#888"}
                _bars = ""
                for _cat, _cnt in _cat_counts.most_common():
                    _pct = int((_cnt / _max_c) * 100) if _max_c else 0
                    _col = _cat_colors.get(_cat, "#a78bfa")
                    _bars += (
                        '<div style="margin-bottom:0.45rem;">'
                        '<div style="display:flex;justify-content:space-between;font-size:0.78rem;color:#aaa;margin-bottom:2px;">'
                        f'<span>{_cat}</span><span>{_cnt}</span></div>'
                        '<div style="background:#1a1a3a;border-radius:4px;height:8px;overflow:hidden;">'
                        f'<div style="width:{_pct}%;height:100%;background:{_col};"></div></div>'
                        '</div>'
                    )
                st.markdown(
                    '<div class="card"><div style="font-size:0.72rem;color:#555;font-family:\'JetBrains Mono\',monospace;letter-spacing:1px;margin-bottom:0.7rem;">THIS MONTH BY CATEGORY · purchases</div>'
                    + _bars + '</div>',
                    unsafe_allow_html=True
                )
            for p in all_purchases:
                currency = p.get("currency", "UZS")
                if not p.get("phase2_completed"):
                    border = "#444"
                    status_badge = '<span class="badge badge-grey">⏳ Pending Review</span>'
                elif p.get("worth_it"):
                    border = "#4ade80"
                    status_badge = '<span class="badge badge-green">✅ Worth It</span>'
                else:
                    border = "#f87171"
                    status_badge = '<span class="badge badge-red">❌ Not Worth It</span>'

                refl_html = f'<div style="margin-top:0.4rem;font-size:0.8rem;color:#666;">💭 {p["review_reflection"]}</div>' if p.get("review_reflection") else ''
                p_card = (
                    f'<div class="card" style="border-left: 3px solid {border};">'
                    '<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.4rem;">'
                    f'<span style="font-weight:800;color:#fff;">{p["item_name"]}</span>'
                    f'<div style="display:flex;gap:0.5rem;">{status_badge}<span class="badge badge-grey">{p.get("category","")}</span></div>'
                    '</div>'
                    '<div style="font-size:0.85rem;color:#888;font-family:\'JetBrains Mono\',monospace;">'
                    f'{p.get("amount",0):,.0f} {currency} · {p.get("purchase_date","")} · {p.get("emotional_state_post_purchase","")}'
                    '</div>'
                    f'{refl_html}'
                    '</div>'
                )
                st.markdown(p_card, unsafe_allow_html=True)

                col_edit, col_del = st.columns([1, 1])
                with col_edit:
                    if st.button("✏️ Edit", key=f"edit_purchase_{p['id']}"):
                        st.session_state["editing_purchase"] = dict(p)
                        st.session_state["purchase_view"] = "➕ Add Purchase"
                        st.rerun()
                with col_del:
                    if st.button("🗑️ Delete", key=f"del_purchase_{p['id']}"):
                        st.session_state[f"confirm_del_p_{p['id']}"] = True
                        st.rerun()
                if st.session_state.get(f"confirm_del_p_{p['id']}", False):
                    st.warning(f"Delete **{p['item_name']}**? This cannot be undone.")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Yes, delete", key=f"yes_del_p_{p['id']}"):
                            supabase.table("purchases").delete().eq("id", p["id"]).execute()
                            st.session_state[f"confirm_del_p_{p['id']}"] = False
                            st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_del_p_{p['id']}"):
                            st.session_state[f"confirm_del_p_{p['id']}"] = False
                            st.rerun()

    # ============================================================
    # ADD / EDIT PURCHASE
    # ============================================================
    else:
        editing_p = st.session_state.get("editing_purchase", None)
        editing_p_id = editing_p["id"] if editing_p else None
        st.markdown(f"#### {'✏️ Editing Purchase' if editing_p else 'What did you buy?'}")

        with st.form("purchase_form"):
            p_name     = st.text_input("Item Name", value=editing_p["item_name"] if editing_p else "", placeholder="e.g. AirPods, lunch, new jacket...")
            categories = ["food", "clothing", "tech", "entertainment", "transport", "other"]
            p_category = st.selectbox("Category", categories, index=categories.index(editing_p["category"]) if editing_p and editing_p.get("category") in categories else 0)
            currencies = ["UZS", "USD", "EUR", "RUB", "GBP", "JPY"]
            col_amt, col_cur = st.columns([3, 1])
            with col_amt:
                p_amount = st.number_input("Amount", min_value=0.0, step=1000.0, value=float(editing_p["amount"]) if editing_p and editing_p.get("amount") else 0.0)
            with col_cur:
                cur_index = currencies.index(editing_p.get("currency", "UZS")) if editing_p and editing_p.get("currency") in currencies else 0
                p_currency = st.selectbox("Currency", currencies, index=cur_index)
            p_reason   = st.text_input("Reason for buying", value=editing_p.get("reason_for_purchase", "") if editing_p else "", placeholder="Why did you buy this?")
            emotions   = ["planned", "impulse", "influenced", "necessity"]
            p_emotion  = st.selectbox("Emotional state at purchase", emotions, index=emotions.index(editing_p["emotional_state_post_purchase"]) if editing_p and editing_p.get("emotional_state_post_purchase") in emotions else 0)
            p_expected = st.text_input("Expected value", value=editing_p.get("expected_value", "") if editing_p else "", placeholder="What do you expect to get from this?")
            from datetime import datetime, timedelta
            default_date = datetime.strptime(editing_p["purchase_date"], "%Y-%m-%d").date() if editing_p and editing_p.get("purchase_date") else date.today()
            p_date     = st.date_input("Purchase Date", value=default_date)
            submitted  = st.form_submit_button("💾 Update Purchase" if editing_p else "💾 Log Purchase")

        if submitted:
            if not p_name.strip():
                st.error("Please enter an item name.")
            else:
                try:
                    from datetime import timedelta
                    record = {
                        "item_name": p_name.strip(),
                        "category": p_category,
                        "amount": p_amount,
                        "currency": p_currency,
                        "reason_for_purchase": p_reason.strip(),
                        "emotional_state_post_purchase": p_emotion,
                        "expected_value": p_expected.strip(),
                        "purchase_date": str(p_date),
                        "review_due_date": str(p_date + timedelta(days=14)),
                    }
                    if editing_p_id:
                        supabase.table("purchases").update(record).eq("id", editing_p_id).execute()
                        st.success("✅ Purchase updated.")
                        st.session_state["editing_purchase"] = None
                    else:
                        record["phase2_completed"] = False
                        supabase.table("purchases").insert(record).execute()
                        st.success(f"✅ Purchase logged. Review set for {p_date + timedelta(days=14)}.")
                    st.session_state["purchase_view"] = "📋 All Purchases"
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

        if editing_p_id and st.button("Cancel", key="cancel_purchase_form"):
            st.session_state["editing_purchase"] = None
            st.session_state["purchase_view"] = "📋 All Purchases"
            st.rerun()
elif page == "🎯  Goals":
    st.markdown('<div class="section-header">🎯 Goals</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">a goal is a dream with a deadline</div>', unsafe_allow_html=True)

    if "goal_view" not in st.session_state:
        st.session_state["goal_view"] = "📋 Goals"
    if "goal_radio_counter" not in st.session_state:
        st.session_state["goal_radio_counter"] = 0

    _g_options = ["📋 Goals", "➕ Add Goal"]
    _g_index = _g_options.index(st.session_state["goal_view"]) if st.session_state["goal_view"] in _g_options else 0
    goal_view = st.radio("Goal View", _g_options, index=_g_index,
                         horizontal=True, label_visibility="collapsed",
                         key=f"goal_radio_{st.session_state['goal_radio_counter']}")
    st.session_state["goal_view"] = goal_view
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # ADD GOAL FORM
    # ============================================================
    if goal_view == "➕ Add Goal":
        st.markdown("#### What's your next goal?")
        _g_cats = ["Personal Growth", "Career", "Academic", "Health & Fitness",
                   "Skills", "Creative", "Financial", "Relationships", "Other"]
        with st.form("add_goal_form"):
            ag_title    = st.text_input("Goal Title", placeholder="e.g. Get IELTS 7.5, Run a marathon, Build an app...")
            ag_col1, ag_col2 = st.columns(2)
            with ag_col1:
                ag_cat  = st.selectbox("Category", _g_cats)
            with ag_col2:
                ag_date = st.date_input("Target Date", value=date.today())
            ag_why      = st.text_input("Why does this matter to you?", placeholder="Your motivation, your reason...")
            ag_notes    = st.text_area("Progress Notes (optional)", height=70, placeholder="Starting point, current state, first steps...")
            ag_submit   = st.form_submit_button("🎯 Create Goal", type="primary")

        if ag_submit:
            if not ag_title.strip():
                st.error("Please enter a goal title.")
            else:
                try:
                    supabase.table("goals").insert({
                        "goal_title": ag_title.strip(),
                        "category": ag_cat,
                        "target_date": str(ag_date),
                        "why": ag_why.strip(),
                        "progress_notes": ag_notes.strip(),
                        "status": "active",
                    }).execute()
                    st.success(f"✅ Goal created: {ag_title.strip()}")
                    st.session_state["goal_view"] = "📋 Goals"
                    st.session_state["goal_radio_counter"] += 1
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

    # ============================================================
    # GOALS LIST
    # ============================================================
    else:
        try:
            goals = supabase.table("goals").select("*").order("created_at", desc=False).execute().data or []
        except Exception:
            goals = []

        active_goals    = [g for g in goals if g["status"] == "active"]
        paused_goals    = [g for g in goals if g["status"] == "paused"]
        completed_goals = [g for g in goals if g["status"] == "completed"]

        if not goals:
            st.info("No goals yet. Create one above or activate a wish from the Wish List.")
        else:
            for g in active_goals + paused_goals + completed_goals:
                try:
                    sessions = supabase.table("goal_sessions").select("*").eq("goal_id", g["id"]).order("session_date", desc=False).execute().data or []
                except Exception:
                    sessions = []

                total_hours   = sum(s.get("duration_hours", 0) or 0 for s in sessions)
                avg_enjoyment = (sum(s.get("enjoyment_score", 0) or 0 for s in sessions) / len(sessions)) if sessions else 0

                # Deadline badge
                deadline_badge = ""
                if g.get("target_date") and g["status"] == "active":
                    try:
                        from datetime import datetime as _dt
                        _td   = _dt.strptime(g["target_date"], "%Y-%m-%d").date()
                        _diff = (_td - date.today()).days
                        if _diff < 0:
                            deadline_badge = f'<span class="badge badge-red">⚠️ {abs(_diff)}d overdue</span>'
                        elif _diff <= 7:
                            deadline_badge = f'<span class="badge badge-yellow">⏰ {_diff}d left</span>'
                        else:
                            deadline_badge = f'<span class="badge badge-blue">📅 {_diff}d left</span>'
                    except Exception:
                        pass

                cat_badge = f'<span class="badge badge-grey">{g["category"]}</span>' if g.get("category") else ""
                # Build as ONE continuous HTML string — blank lines break Streamlit's HTML renderer
                why_html   = f'<div style="font-size:0.82rem;color:#a78bfa;margin-bottom:0.3rem;">💡 {g["why"]}</div>' if g.get("why") else ""
                notes_html = f'<div style="font-size:0.82rem;color:#666;margin-bottom:0.3rem;">📝 {g["progress_notes"]}</div>' if g.get("progress_notes") else ""
                status_html = goal_status_badge(g["status"])

                card_html = (
                    '<div class="card">'
                    '<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">'
                    f'<span style="font-size:1.1rem;font-weight:800;color:#fff;">{g["goal_title"]}</span>'
                    f'<div style="display:flex;gap:0.5rem;align-items:center;">{cat_badge} {deadline_badge} {status_html}</div>'
                    '</div>'
                    f'{why_html}{notes_html}'
                    '<div style="font-size:0.85rem;color:#888;font-family:\'JetBrains Mono\',monospace;">'
                    f'{len(sessions)} sessions · {total_hours:.1f} hrs total · avg enjoyment {avg_enjoyment:.1f}/5'
                    '</div>'
                    '</div>'
                )
                st.markdown(card_html, unsafe_allow_html=True)

                if sessions:
                    with st.expander(f"🗂 {len(sessions)} past sessions"):
                        for s in reversed(sessions):
                            stars = '★' * (s.get('enjoyment_score') or 0) + '☆' * (5 - (s.get('enjoyment_score') or 0))
                            s_desc    = s.get('activity_description','')
                            s_summary = s.get('learning_summary','')
                            st.markdown(f"""
                            <div class="card" style="margin-bottom:0.4rem;">
                                <div style="display:flex;justify-content:space-between;">
                                    <span style="font-family:'JetBrains Mono',monospace;font-size:0.8rem;color:#888;">{s['session_date']}</span>
                                    <span class="badge badge-blue">{s.get('duration_hours',0)} hrs</span>
                                    <span class="badge badge-yellow">{stars}</span>
                                </div>
                                <div style="margin-top:0.4rem;color:#ccc;font-size:0.85rem;"><b>{s_desc}</b></div>
                                <div style="margin-top:0.2rem;color:#888;font-size:0.85rem;">{s_summary}</div>
                            </div>
                            """, unsafe_allow_html=True)

                if g["status"] == "active":
                    with st.expander(f"➕ Log a session for '{g['goal_title']}'"):
                        s_date     = st.date_input("Session Date", value=date.today(), key=f"sdate_{g['id']}")
                        s_duration = st.number_input("Duration (hours)", min_value=0.0, step=0.5, value=1.0, key=f"sdur_{g['id']}")
                        s_activity = st.text_input("What did you do?", key=f"sact_{g['id']}", placeholder="Describe the activity...")
                        s_learning = st.text_area("What did you learn?", key=f"slearn_{g['id']}", height=80)
                        s_enjoy    = st.slider("Enjoyment (1–5)", 1, 5, 3, key=f"senjoy_{g['id']}")

                        if st.button("💾 Save Session", key=f"save_goal_session_{g['id']}", type="primary"):
                            try:
                                supabase.table("goal_sessions").insert({
                                    "goal_id": g["id"],
                                    "session_date": str(s_date),
                                    "duration_hours": s_duration,
                                    "activity_description": s_activity.strip(),
                                    "learning_summary": s_learning.strip(),
                                    "enjoyment_score": s_enjoy
                                }).execute()
                                st.success("✅ Session logged.")
                                st.rerun()
                            except Exception as ex:
                                st.error(f"Error: {ex}")

                if g["status"] == "active":
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("⏸️ Pause", key=f"pause_{g['id']}"):
                            supabase.table("goals").update({"status": "paused"}).eq("id", g["id"]).execute()
                            st.rerun()
                    with col2:
                        if st.button("🏁 End → Outcome", key=f"end_{g['id']}", type="primary"):
                            st.session_state["outcome_goal"] = dict(g)
                            st.session_state["current_page_override"] = "🏆  Outcomes"
                            st.session_state["nav_override_counter"] = st.session_state.get("nav_override_counter", 0) + 1
                            st.rerun()
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_goal_{g['id']}"):
                            st.session_state[f"confirm_del_g_{g['id']}"] = True
                            st.rerun()
                elif g["status"] == "paused":
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("▶️ Resume Goal", key=f"resume_{g['id']}"):
                            supabase.table("goals").update({"status": "active"}).eq("id", g["id"]).execute()
                            st.rerun()
                    with col2:
                        if st.button("🗑️ Delete", key=f"del_goal_{g['id']}"):
                            st.session_state[f"confirm_del_g_{g['id']}"] = True
                            st.rerun()
                elif g["status"] == "completed":
                    if st.button("🗑️ Delete", key=f"del_goal_{g['id']}"):
                        st.session_state[f"confirm_del_g_{g['id']}"] = True
                        st.rerun()

                if st.session_state.get(f"confirm_del_g_{g['id']}", False):
                    st.warning(f"Delete **{g['goal_title']}** and all its sessions?")
                    _dc1, _dc2 = st.columns(2)
                    with _dc1:
                        if st.button("Yes, delete", key=f"yes_del_g_{g['id']}"):
                            supabase.table("goal_sessions").delete().eq("goal_id", g["id"]).execute()
                            supabase.table("goals").delete().eq("id", g["id"]).execute()
                            st.session_state[f"confirm_del_g_{g['id']}"] = False
                            st.rerun()
                    with _dc2:
                        if st.button("Cancel", key=f"cancel_del_g_{g['id']}"):
                            st.session_state[f"confirm_del_g_{g['id']}"] = False
                            st.rerun()

                st.markdown('<hr class="divider">', unsafe_allow_html=True)

elif page == "💫  Wish List":
    st.markdown('<div class="section-header">💫 Wish List</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">no deadlines · no pressure · just things you want</div>', unsafe_allow_html=True)

    if "wish_view" not in st.session_state:
        st.session_state["wish_view"] = "💫 Passive Wishes"
    if "editing_wish" not in st.session_state:
        st.session_state["editing_wish"] = None
    if "wish_radio_counter" not in st.session_state:
        st.session_state["wish_radio_counter"] = 0

    options = ["💫 Passive Wishes", "➕ Add Wish"]
    current_index = options.index(st.session_state["wish_view"])
    view = st.radio("Wish View", options, index=current_index,
                    horizontal=True, label_visibility="collapsed",
                    key=f"wish_radio_{st.session_state['wish_radio_counter']}")
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # PASSIVE WISHES
    # ============================================================
    if view == "💫 Passive Wishes":
        try:
            wishes = supabase.table("wishes").select("*").eq("status", "passive").order("created_at", desc=True).execute().data or []
        except Exception:
            wishes = []

        if not wishes:
            st.info("No wishes yet. Add anything you want to pursue someday.")
        else:
            # Filter bar
            _pf_opts = ["All", "🔴 High", "🟡 Medium", "🟢 Low"]
            _pf = st.radio("Filter by priority", _pf_opts, horizontal=True,
                           label_visibility="collapsed", key="wish_priority_filter")
            _pf_map = {"All": None, "🔴 High": "high", "🟡 Medium": "medium", "🟢 Low": "low"}
            _pf_val = _pf_map[_pf]

            # Sort: high → medium → low → None
            _priority_order = {"high": 0, "medium": 1, "low": 2, None: 3, "": 3}
            wishes_sorted = sorted(wishes, key=lambda w: _priority_order.get(w.get("priority",""), 3))
            if _pf_val:
                wishes_sorted = [w for w in wishes_sorted if w.get("priority") == _pf_val]

            _priority_colors = {"high": "#f87171", "medium": "#facc15", "low": "#4ade80"}
            _priority_labels = {"high": "🔴 HIGH", "medium": "🟡 MEDIUM", "low": "🟢 LOW"}

            for w in wishes_sorted:
                pri = w.get("priority", "")
                pri_color = _priority_colors.get(pri, "#555")
                pri_label = _priority_labels.get(pri, "")
                cat_badge = f'<span class="badge badge-grey">{w["category"]}</span>' if w.get("category") else ""
                pri_badge = f'<span style="font-size:0.72rem;font-weight:700;color:{pri_color};font-family:JetBrains Mono,monospace;">{pri_label}</span>' if pri_label else ""

                # "added X days ago"
                try:
                    from datetime import datetime as _dt
                    _added = _dt.strptime(w.get("date_added",""), "%Y-%m-%d").date()
                    _days_ago = (date.today() - _added).days
                    _ago_str = "today" if _days_ago == 0 else f"{_days_ago} days ago"
                except Exception:
                    _ago_str = w.get("date_added","")

                st.markdown(f"""
                <div class="card">
                    <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:0.3rem;">
                        <div style="font-size:1rem;font-weight:800;color:#fff;">✦ {w['wish_title']}</div>
                        <div style="display:flex;gap:0.5rem;align-items:center;">{pri_badge} {cat_badge}</div>
                    </div>
                    <div style="color:#aaa;font-size:0.9rem;">{w.get('description','')}</div>
                    <div style="margin-top:0.4rem;font-size:0.75rem;color:#555;font-family:'JetBrains Mono',monospace;">added {_ago_str}</div>
                </div>
                """, unsafe_allow_html=True)

                col1, col2, col3 = st.columns([1, 1, 1])
                with col1:
                    if st.button("🎯 Activate as Goal", key=f"activate_{w['id']}", type="primary"):
                        try:
                            supabase.table("wishes").update({"status": "activated"}).eq("id", w["id"]).execute()
                            supabase.table("goals").insert({
                                "goal_title": w["wish_title"],
                                "linked_wish_id": w["id"],
                                "status": "active",
                                "category": w.get("category", ""),
                                "why": w.get("description", ""),
                            }).execute()
                            st.success(f"✅ '{w['wish_title']}' is now an active goal — category and motivation carried over!")
                            st.rerun()
                        except Exception as ex:
                            st.error(f"Error: {ex}")
                with col2:
                    if st.button("✏️ Edit", key=f"edit_wish_{w['id']}"):
                        st.session_state["editing_wish"] = dict(w)
                        st.session_state["wish_view"] = "➕ Add Wish"
                        st.session_state["wish_radio_counter"] += 1
                        st.rerun()
                with col3:
                    if st.button("🗑️ Delete", key=f"del_wish_{w['id']}"):
                        st.session_state[f"confirm_del_w_{w['id']}"] = True
                        st.rerun()

                if st.session_state.get(f"confirm_del_w_{w['id']}", False):
                    st.warning(f"Delete **{w['wish_title']}**? This cannot be undone.")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Yes, delete", key=f"yes_del_w_{w['id']}"):
                            supabase.table("wishes").delete().eq("id", w["id"]).execute()
                            st.session_state[f"confirm_del_w_{w['id']}"] = False
                            st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_del_w_{w['id']}"):
                            st.session_state[f"confirm_del_w_{w['id']}"] = False
                            st.rerun()

                st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # ADD / EDIT WISH
    # ============================================================
    else:
        editing_w = st.session_state.get("editing_wish", None)
        editing_w_id = editing_w["id"] if editing_w else None
        st.markdown(f"#### {'✏️ Editing Wish' if editing_w else 'What do you want?'}")

        with st.form("wish_form"):
            w_title = st.text_input("Wish Title", value=editing_w["wish_title"] if editing_w else "", placeholder="e.g. Learn Japanese, Visit Japan, Build a startup...")
            w_desc  = st.text_area("Description (optional)", value=editing_w.get("description","") if editing_w else "", height=100, placeholder="Any details, context, or why this matters to you...")
            _w_cats = ["Personal Growth", "Travel", "Career", "Skills", "Health & Fitness",
                       "Relationships", "Creative", "Financial", "Education", "Long-term Dream", "Other"]
            _cat_default = _w_cats.index(editing_w.get("category","Other")) if editing_w and editing_w.get("category") in _w_cats else len(_w_cats)-1
            w_col1, w_col2 = st.columns(2)
            with w_col1:
                w_category = st.selectbox("Category", _w_cats, index=_cat_default)
            with w_col2:
                _pri_opts = ["high", "medium", "low"]
                _pri_default = _pri_opts.index(editing_w.get("priority","medium")) if editing_w and editing_w.get("priority") in _pri_opts else 1
                w_priority = st.selectbox("Priority", _pri_opts, index=_pri_default)
            submitted = st.form_submit_button("💾 Update Wish" if editing_w else "💾 Add Wish")

        if submitted:
            if not w_title.strip():
                st.error("Please enter a wish title.")
            else:
                try:
                    record = {
                        "wish_title": w_title.strip(),
                        "description": w_desc.strip(),
                        "category": w_category,
                        "priority": w_priority,
                    }
                    if editing_w_id:
                        supabase.table("wishes").update(record).eq("id", editing_w_id).execute()
                        st.success("✅ Wish updated.")
                        st.session_state["editing_wish"] = None
                    else:
                        record["status"] = "passive"
                        record["date_added"] = str(date.today())
                        supabase.table("wishes").insert(record).execute()
                        st.success("✅ Wish added to your list.")
                    st.session_state["wish_view"] = "💫 Passive Wishes"
                    st.session_state["wish_radio_counter"] += 1
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

        if editing_w_id and st.button("Cancel", key="cancel_wish_form"):
            st.session_state["editing_wish"] = None
            st.session_state["wish_view"] = "💫 Passive Wishes"
            st.session_state["wish_radio_counter"] += 1
            st.rerun()

elif page == "🏆  Outcomes":
    st.markdown('<div class="section-header">🏆 Outcomes</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">where goals end and become permanent record</div>', unsafe_allow_html=True)

    if "outcome_view" not in st.session_state:
        st.session_state["outcome_view"] = "📋 All Outcomes"
    if "editing_outcome" not in st.session_state:
        st.session_state["editing_outcome"] = None
    if "outcome_goal" not in st.session_state:
        st.session_state["outcome_goal"] = None

    # If a goal was passed in to end (from Goals page), force Log view
    if st.session_state.get("outcome_goal"):
        st.session_state["outcome_view"] = "➕ Log Outcome"

    options = ["📋 All Outcomes", "➕ Log Outcome"]
    current_index = options.index(st.session_state["outcome_view"])
    view = st.radio("Outcome View", options, index=current_index,
                    horizontal=True, label_visibility="collapsed", key=f"outcome_radio_{current_index}")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # ALL OUTCOMES
    # ============================================================
    if view == "📋 All Outcomes":
        _result_opts = {"All": None, "✅ Win": "win", "❌ Fail": "fail", "🔵 Pass": "pass", "🏁 Complete": "complete"}
        filter_label = st.selectbox("Filter by result", list(_result_opts.keys()), label_visibility="collapsed")
        filter_type = _result_opts[filter_label]

        try:
            query = supabase.table("outcomes").select("*").order("outcome_date", desc=True)
            if filter_type:
                query = query.eq("result_type", filter_type)
            outcomes = query.execute().data or []
        except Exception:
            outcomes = []

        if not outcomes:
            st.info("No outcomes logged yet. End a goal from the Goals page to record one.")
        else:
            border_colors = {"win": "#4ade80", "pass": "#60a5fa", "fail": "#f87171", "complete": "#fbbf24"}
            for o in outcomes:
                border = border_colors.get(o.get("result_type"), "#444")
                st.markdown(f"""
                <div class="card" style="border-left: 3px solid {border};">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
                        <span style="font-size:1rem;font-weight:800;color:#fff;">{o['outcome_title']}</span>
                        {result_badge(o.get('result_type','pass'))}
                    </div>
                    <div style="color:#ccc;font-size:0.9rem;margin-bottom:0.4rem;"><b>Result:</b> {o.get('score_or_result','')}</div>
                    <div style="color:#888;font-size:0.85rem;font-style:italic;margin-bottom:0.4rem;">💭 {o.get('emotional_reaction','')}</div>
                    <div style="color:#888;font-size:0.85rem;margin-bottom:0.4rem;"><b>Why:</b> {o.get('causal_analysis','')}</div>
                    <div style="color:#aaa;font-size:0.85rem;"><b>💡 Lesson:</b> {o.get('lessons_learned','')}</div>
                    <div style="margin-top:0.5rem;font-size:0.75rem;color:#555;font-family:'JetBrains Mono',monospace;">{o.get('outcome_date','')}</div>
                </div>
                """, unsafe_allow_html=True)

                col_edit, col_share, col_del = st.columns([1, 1, 1])
                with col_edit:
                    if st.button("✏️ Edit", key=f"edit_outcome_{o['id']}"):
                        st.session_state["editing_outcome"] = dict(o)
                        st.session_state["outcome_view"] = "➕ Log Outcome"
                        st.rerun()
                with col_share:
                    if st.button("📤 Share", key=f"share_outcome_{o['id']}"):
                        st.session_state[f"show_share_o_{o['id']}"] = not st.session_state.get(f"show_share_o_{o['id']}", False)
                        st.rerun()
                with col_del:
                    if st.button("🗑️ Delete", key=f"del_outcome_{o['id']}"):
                        st.session_state[f"confirm_del_o_{o['id']}"] = True
                        st.rerun()

                if st.session_state.get(f"show_share_o_{o['id']}", False):
                    _result_emoji = {"win":"✅","fail":"❌","pass":"🔵","complete":"🏁"}.get(o.get("result_type",""),"🏆")
                    _share_lines = [
                        f"🏆 Life Archive — Outcome",
                        f"📅 {o.get('outcome_date','')}",
                        f"",
                        f"{o.get('outcome_title','')}",
                        f"Result: {_result_emoji} {(o.get('result_type') or '').upper()}",
                        f"",
                        f"Summary: {o.get('score_or_result','')}",
                        f"💭 {o.get('emotional_reaction','')}",
                        f"💡 Lesson: {o.get('lessons_learned','')}",
                        f"",
                        f"— sent from Life Archive",
                    ]
                    _share_text = "\n".join(_share_lines)
                    st.text_area("Share this outcome", value=_share_text, height=220,
                                 key=f"share_text_o_{o['id']}")
                    _tg_o = f"https://t.me/share/url?url=&text={_share_text.replace(' ','%20').replace(chr(10),'%0A')[:2000]}"
                    st.markdown(f'<a href="{_tg_o}" target="_blank" style="display:inline-block;margin-top:0.3rem;padding:0.4rem 1.2rem;background:#2a2a4a;border:1px solid #a78bfa;border-radius:8px;color:#e8e0ff;text-decoration:none;font-family:JetBrains Mono,monospace;font-size:0.8rem;">📨 Open in Telegram</a>', unsafe_allow_html=True)
                if st.session_state.get(f"confirm_del_o_{o['id']}", False):
                    st.warning(f"Delete **{o['outcome_title']}**? This cannot be undone.")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Yes, delete", key=f"yes_del_o_{o['id']}"):
                            supabase.table("outcomes").delete().eq("id", o["id"]).execute()
                            supabase.table("win_failure_archive").delete().eq("outcome_id", o["id"]).execute()
                            st.session_state[f"confirm_del_o_{o['id']}"] = False
                            st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_del_o_{o['id']}"):
                            st.session_state[f"confirm_del_o_{o['id']}"] = False
                            st.rerun()

    # ============================================================
    # LOG / EDIT OUTCOME
    # ============================================================
    else:
        editing_o = st.session_state.get("editing_outcome", None)
        editing_o_id = editing_o["id"] if editing_o else None
        goal_ctx = st.session_state.get("outcome_goal", None)

        if editing_o:
            st.markdown("#### ✏️ Editing Outcome")
        elif goal_ctx:
            st.markdown(f"#### 🏁 Ending Goal: *{goal_ctx['goal_title']}*")
        else:
            st.markdown("#### Log a New Outcome")

        with st.form("outcome_form"):
            default_title = editing_o["outcome_title"] if editing_o else (goal_ctx["goal_title"] if goal_ctx else "")
            o_title = st.text_input("Outcome Title", value=default_title, placeholder="e.g. IELTS attempt, Startup pitch...")

            result_types = ["win", "fail", "pass", "complete"]
            default_result = editing_o["result_type"] if editing_o and editing_o.get("result_type") in result_types else "complete"
            o_result_type = st.selectbox("Result Type", result_types, index=result_types.index(default_result))

            o_score = st.text_input("Score / Result", value=editing_o.get("score_or_result","") if editing_o else "", placeholder="e.g. Scored 8.5, Got the job, Didn't finish...")
            o_emotion = st.text_area("Emotional Reaction", value=editing_o.get("emotional_reaction","") if editing_o else "", height=80, placeholder="How did you feel about it?")
            o_env = st.text_area("Environment Context", value=editing_o.get("environment_context","") if editing_o else "", height=80, placeholder="What was going on around you at the time?")
            o_causal = st.text_area("Causal Analysis", value=editing_o.get("causal_analysis","") if editing_o else "", height=80, placeholder="Why did it go this way?")
            o_lessons = st.text_area("Lessons Learned", value=editing_o.get("lessons_learned","") if editing_o else "", height=80, placeholder="What will you take from this?")

            from datetime import datetime
            default_date = datetime.strptime(editing_o["outcome_date"], "%Y-%m-%d").date() if editing_o and editing_o.get("outcome_date") else date.today()
            o_date = st.date_input("Outcome Date", value=default_date)

            submitted = st.form_submit_button("💾 Update Outcome" if editing_o else "💾 Save Outcome")

        if submitted:
            if not o_title.strip():
                st.error("Please enter an outcome title.")
            else:
                try:
                    record = {
                        "outcome_title": o_title.strip(),
                        "result_type": o_result_type,
                        "score_or_result": o_score.strip(),
                        "emotional_reaction": o_emotion.strip(),
                        "environment_context": o_env.strip(),
                        "causal_analysis": o_causal.strip(),
                        "lessons_learned": o_lessons.strip(),
                        "outcome_date": str(o_date),
                    }

                    if editing_o_id:
                        supabase.table("outcomes").update(record).eq("id", editing_o_id).execute()
                        # Update archive copy too
                        supabase.table("win_failure_archive").update({
                            "title": o_title.strip(),
                            "result_type": o_result_type,
                            "summary": o_score.strip(),
                            "reflection": o_lessons.strip(),
                            "archive_date": str(o_date)
                        }).eq("outcome_id", editing_o_id).execute()
                        st.success("✅ Outcome updated.")
                        st.session_state["editing_outcome"] = None
                    else:
                        if goal_ctx:
                            record["linked_goal_id"] = goal_ctx["id"]

                        result = supabase.table("outcomes").insert(record).execute()
                        new_outcome_id = result.data[0]["id"]

                        # Push to win/failure archive
                        supabase.table("win_failure_archive").insert({
                            "outcome_id": new_outcome_id,
                            "archive_date": str(o_date),
                            "title": o_title.strip(),
                            "result_type": o_result_type,
                            "summary": o_score.strip(),
                            "reflection": o_lessons.strip()
                        }).execute()

                        # Mark linked goal as completed
                        if goal_ctx:
                            supabase.table("goals").update({"status": "completed"}).eq("id", goal_ctx["id"]).execute()
                            st.session_state["outcome_goal"] = None

                        st.success("✅ Outcome saved and archived.")

                    st.session_state["outcome_view"] = "📋 All Outcomes"
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

        if (editing_o_id or goal_ctx) and st.button("Cancel", key="cancel_outcome_form"):
            st.session_state["editing_outcome"] = None
            st.session_state["outcome_goal"] = None
            st.session_state["outcome_view"] = "📋 All Outcomes"
            st.rerun()

elif page == "📜  Archive":
    st.markdown('<div class="section-header">📜 Archive</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">your life, browsable</div>', unsafe_allow_html=True)

    _arch_opts = ["🏆 Win/Failure Archive", "📅 Daily Log History", "📊 Quick Stats"]
    if "archive_view" not in st.session_state or st.session_state["archive_view"] not in _arch_opts:
        st.session_state["archive_view"] = "🏆 Win/Failure Archive"
    _arch_index = _arch_opts.index(st.session_state["archive_view"])
    archive_view = st.radio("Archive View", _arch_opts, index=_arch_index,
                            horizontal=True, label_visibility="collapsed",
                            key=f"archive_radio_{_arch_index}")
    st.session_state["archive_view"] = archive_view
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ============================================================
    # WIN/FAILURE ARCHIVE
    # ============================================================
    if archive_view == "🏆 Win/Failure Archive":
        filter_type = st.selectbox("Filter", ["All", "win", "fail", "pass", "complete"], label_visibility="collapsed")

        try:
            query = supabase.table("win_failure_archive").select("*").order("archive_date", desc=True)
            if filter_type != "All":
                query = query.eq("result_type", filter_type)
            entries = query.execute().data or []
        except Exception:
            entries = []

        if not entries:
            st.info("Your win/failure archive is empty. Outcomes you log will appear here permanently.")
        else:
            border_colors = {"win": "#4ade80", "pass": "#60a5fa", "fail": "#f87171", "complete": "#fbbf24"}
            for e in entries:
                border = border_colors.get(e.get("result_type"), "#444")
                st.markdown(f"""
                <div class="card" style="border-left: 3px solid {border};">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.4rem;">
                        <span style="font-size:1rem;font-weight:800;color:#fff;">{e['title']}</span>
                        {result_badge(e.get('result_type','pass'))}
                    </div>
                    <div style="color:#ccc;font-size:0.9rem;margin-bottom:0.4rem;">{e.get('summary','')}</div>
                    <div style="color:#888;font-size:0.85rem;font-style:italic;">💡 {e.get('reflection','')}</div>
                    <div style="margin-top:0.5rem;font-size:0.75rem;color:#555;font-family:'JetBrains Mono',monospace;">{e.get('archive_date','')}</div>
                </div>
                """, unsafe_allow_html=True)

    # ============================================================
    # DAILY LOG HISTORY
    # ============================================================
    elif archive_view == "📅 Daily Log History":
        try:
            logs = supabase.table("daily_logs").select("*").order("date", desc=True).execute().data or []
        except Exception:
            logs = []

        if not logs:
            st.info("No daily logs yet. Start with today's entry.")
        else:
            # Mood trend chart
            chart_logs = list(reversed(logs))
            dates = [l["date"] for l in chart_logs]
            moods = [l.get("mood_score", 5) for l in chart_logs]

            svg_width = 800
            svg_height = 150
            padding = 20
            n = len(moods)
            if n > 1:
                x_step = (svg_width - 2 * padding) / (n - 1)
                points = []
                for i, m in enumerate(moods):
                    x = padding + i * x_step
                    y = svg_height - padding - ((m - 1) / 9) * (svg_height - 2 * padding)
                    points.append(f"{x:.1f},{y:.1f}")
                points_str = " ".join(points)

                circles = ""
                for i, m in enumerate(moods):
                    x = padding + i * x_step
                    y = svg_height - padding - ((m - 1) / 9) * (svg_height - 2 * padding)
                    color = "#4ade80" if m >= 8 else "#facc15" if m >= 5 else "#f87171"
                    circles += f'<circle cx="{x:.1f}" cy="{y:.1f}" r="3" fill="{color}"/>'

                st.markdown(f"""
                <div class="card">
                    <div style="font-size:0.75rem;color:#555;font-family:'JetBrains Mono',monospace;margin-bottom:0.5rem;">MOOD TREND</div>
                    <svg width="100%" height="{svg_height}" viewBox="0 0 {svg_width} {svg_height}" preserveAspectRatio="none">
                        <polyline points="{points_str}" fill="none" stroke="#60a5fa" stroke-width="2"/>
                        {circles}
                    </svg>
                    <div style="display:flex;justify-content:space-between;font-size:0.7rem;color:#555;font-family:'JetBrains Mono',monospace;margin-top:0.3rem;">
                        <span>{dates[0]}</span><span>{dates[-1]}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown('<hr class="divider">', unsafe_allow_html=True)

            # Browse by date
            for log in logs:
                with st.expander(f"{log['date']} — Mood {log.get('mood_score','—')}/10"):
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.markdown(f"**Mood** &nbsp; {mood_badge(log.get('mood_score',5))}", unsafe_allow_html=True)
                    with col2:
                        st.markdown(f"**Clarity** &nbsp; {clarity_badge(log.get('mental_clarity','normal'))}", unsafe_allow_html=True)
                    with col3:
                        st.markdown(f"**Emotion** &nbsp; `{log.get('dominant_emotion','—')}`")

                    st.markdown(f"**Self Assessment:** {log.get('self_assessment','—')}")
                    st.markdown(f"**Daily Summary:** {log.get('daily_summary','—')}")

                    accs = log.get('accomplishments') or []
                    if accs:
                        st.markdown("**Accomplishments:**")
                        for a in accs:
                            st.markdown(f"&nbsp; ✦ {a}")

                    media = log.get('media_consumed') or []
                    if media:
                        st.markdown("**Media:**")
                        for m in media:
                            st.markdown(f"&nbsp; 🎬 {m}")

                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.markdown(f"**Sleep** `{log.get('sleep_duration','—')} hrs`")
                    with col2:
                        st.markdown(f"**Physical** {physical_badge(log.get('physical_state','neutral'))}", unsafe_allow_html=True)
                    with col3:
                        st.markdown(f"**Mental** {mental_badge(log.get('mental_state','stable'))}", unsafe_allow_html=True)
                    with col4:
                        st.markdown(f"**Spent** `{log.get('daily_spending',0):,.0f}`")

                    st.markdown(f"**Good deed:** {log.get('good_deed','—')}")

    # ============================================================
    # QUICK STATS
    # ============================================================
    else:
        try:
            logs = supabase.table("daily_logs").select("*").execute().data or []
            books = supabase.table("books").select("*").execute().data or []
            events = supabase.table("life_events").select("*").execute().data or []
            outcomes = supabase.table("outcomes").select("*").execute().data or []
        except Exception:
            logs = books = events = outcomes = []

        total_days = len(logs)
        avg_mood = (sum(l.get("mood_score", 0) for l in logs) / total_days) if total_days else 0

        from collections import Counter
        emotions = [l.get("dominant_emotion","").strip().lower() for l in logs if l.get("dominant_emotion")]
        most_common_emotion = Counter(emotions).most_common(1)[0][0] if emotions else "—"

        finished_books = [b for b in books if b["status"] == "finished"]

        wins = len([o for o in outcomes if o.get("result_type") == "win"])
        fails = len([o for o in outcomes if o.get("result_type") == "fail"])

        sig_breakdown = Counter(e.get("significance_score", 0) for e in events)

        st.markdown(f"""
        <div class="metric-row">
            <div class="metric-card">
                <div class="metric-value" style="color:#4ade80">{total_days}</div>
                <div class="metric-label">Days Logged</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#60a5fa">{avg_mood:.1f}</div>
                <div class="metric-label">Avg Mood</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#facc15;font-size:1.2rem;">{most_common_emotion}</div>
                <div class="metric-label">Top Emotion</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#fb923c">{len(finished_books)}</div>
                <div class="metric-label">Books Finished</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""
        <div class="metric-row">
            <div class="metric-card">
                <div class="metric-value" style="color:#4ade80">{wins}</div>
                <div class="metric-label">Wins</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#f87171">{fails}</div>
                <div class="metric-label">Fails</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#a78bfa">{len(events)}</div>
                <div class="metric-label">Life Events</div>
            </div>
            <div class="metric-card">
                <div class="metric-value" style="color:#fbbf24">{sig_breakdown.get(5,0)}</div>
                <div class="metric-label">5★ Moments</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("""
        <div class="card">
            <div style="font-size:0.75rem;color:#555;font-family:'JetBrains Mono',monospace;margin-bottom:0.5rem;">REFLECTION</div>
            <div style="color:#ccc;">
                Every number here represents something real you lived through. This archive doesn't judge — it just remembers, so you don't have to carry it all in your head.
            </div>
        </div>
        """, unsafe_allow_html=True)

# ============================================
# SKILLS PAGE
# ============================================
elif page == "🧠  Skills":
    st.markdown('<div class="section-header">🧠 Skills</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">track mastery · hour by hour</div>', unsafe_allow_html=True)

    if "skill_view" not in st.session_state:
        st.session_state["skill_view"] = "📊 Active Skills"
    if "editing_skill" not in st.session_state:
        st.session_state["editing_skill"] = None
    if "skill_radio_counter" not in st.session_state:
        st.session_state["skill_radio_counter"] = 0

    options = ["📊 Active Skills", "➕ Add Skill"]
    current_index = options.index(st.session_state["skill_view"])
    view = st.radio("Skill View", options, index=current_index,
                    horizontal=True, label_visibility="collapsed",
                    key=f"skill_radio_{st.session_state['skill_radio_counter']}")
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── ACTIVE SKILLS ──
    if view == "📊 Active Skills":
        try:
            skills = supabase.table("skills").select("*").order("created_at", desc=False).execute().data or []
        except Exception:
            skills = []

        if not skills:
            st.info("No skills yet. Add your first skill to start tracking mastery.")
        else:
            for sk in skills:
                try:
                    sessions = supabase.table("skill_sessions").select("*").eq("skill_id", sk["id"]).order("session_date", desc=False).execute().data or []
                except Exception:
                    sessions = []

                total_hours = sum(s.get("duration_hours", 0) or 0 for s in sessions)
                target_hours = sk.get("target_hours", 100) or 100
                progress_pct = min(int((total_hours / target_hours) * 100), 100)
                last_date = sessions[-1]["session_date"] if sessions else None

                from datetime import datetime, timedelta
                inactive_badge = ""
                if last_date:
                    days_since = (date.today() - datetime.strptime(last_date, "%Y-%m-%d").date()).days
                    if days_since >= 5:
                        inactive_badge = f'<span class="badge badge-red">🔴 {days_since}d inactive</span>'

                progress_color = "#a78bfa" if progress_pct < 50 else "#818cf8" if progress_pct < 80 else "#4ade80"

                st.markdown(f"""
                <div class="card">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
                        <span style="font-size:1.1rem;font-weight:800;color:#e8e0ff;">{sk['skill_name']}</span>
                        <div style="display:flex;gap:0.5rem;align-items:center;">
                            {skill_level_badge(sk.get('target_level','beginner'))}
                            {inactive_badge}
                        </div>
                    </div>
                    <div style="font-size:0.8rem;color:#3a3a6a;font-family:'JetBrains Mono',monospace;margin-bottom:0.6rem;">
                        {sk.get('category','')} · {total_hours:.1f} / {target_hours} hrs · {len(sessions)} sessions
                    </div>
                    <div class="progress-bar-bg">
                        <div class="progress-bar-fill" style="width:{progress_pct}%;background:{progress_color};"></div>
                    </div>
                    <div style="font-size:0.75rem;color:#3a3a6a;font-family:'JetBrains Mono',monospace;margin-top:3px;">{progress_pct}% to {sk.get('target_level','goal')}</div>
                </div>
                """, unsafe_allow_html=True)

                if sessions:
                    with st.expander(f"🗂 {len(sessions)} past sessions"):
                        for s in reversed(sessions):
                            stars = '★' * (s.get('difficulty',0) or 0) + '☆' * (5 - (s.get('difficulty',0) or 0))
                            st.markdown(f"""
                            <div class="card" style="margin-bottom:0.4rem;">
                                <div style="display:flex;justify-content:space-between;">
                                    <span style="font-family:'JetBrains Mono',monospace;font-size:0.8rem;color:#3a3a6a;">{s['session_date']}</span>
                                    <span class="badge badge-purple">{s.get('duration_hours',0)}h</span>
                                    <span class="badge badge-yellow">{stars}</span>
                                </div>
                                <div style="margin-top:0.4rem;color:#e8e0ff;font-size:0.85rem;"><b>{s.get('what_practiced','')}</b></div>
                                <div style="margin-top:0.2rem;color:#6a6a9a;font-size:0.85rem;">💡 {s.get('learning_note','')}</div>
                            </div>
                            """, unsafe_allow_html=True)

                with st.expander(f"➕ Log a session for '{sk['skill_name']}'"):
                    s_date     = st.date_input("Date", value=date.today(), key=f"sk_date_{sk['id']}")
                    s_dur      = st.number_input("Duration (hours)", min_value=0.0, step=0.25, value=1.0, key=f"sk_dur_{sk['id']}")
                    s_what     = st.text_input("What did you practice?", key=f"sk_what_{sk['id']}", placeholder="e.g. Hiragana drills, jab-cross combos, recursion problems...")
                    s_note     = st.text_area("What clicked / what you learned", key=f"sk_note_{sk['id']}", height=70)
                    s_diff     = st.slider("Difficulty (1–5)", 1, 5, 3, key=f"sk_diff_{sk['id']}")
                    s_enjoy    = st.slider("Enjoyment (1–5)", 1, 5, 3, key=f"sk_enjoy_{sk['id']}")

                    if st.button("💾 Save Session", key=f"sk_save_{sk['id']}", type="primary"):
                        try:
                            supabase.table("skill_sessions").insert({
                                "skill_id": sk["id"],
                                "session_date": str(s_date),
                                "duration_hours": s_dur,
                                "what_practiced": s_what.strip(),
                                "learning_note": s_note.strip(),
                                "difficulty": s_diff,
                                "enjoyment": s_enjoy
                            }).execute()
                            st.success("✅ Session logged.")
                            st.rerun()
                        except Exception as ex:
                            st.error(f"Error: {ex}")

                col_edit, col_del = st.columns([1, 1])
                with col_edit:
                    if st.button("✏️ Edit", key=f"edit_skill_{sk['id']}"):
                        st.session_state["editing_skill"] = dict(sk)
                        st.session_state["skill_view"] = "➕ Add Skill"
                        st.session_state["skill_radio_counter"] += 1
                        st.rerun()
                with col_del:
                    if st.button("🗑️ Delete", key=f"del_skill_{sk['id']}"):
                        st.session_state[f"confirm_del_sk_{sk['id']}"] = True
                        st.rerun()
                if st.session_state.get(f"confirm_del_sk_{sk['id']}", False):
                    st.warning(f"Delete **{sk['skill_name']}**? All sessions will be lost.")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Yes, delete", key=f"yes_del_sk_{sk['id']}"):
                            supabase.table("skill_sessions").delete().eq("skill_id", sk["id"]).execute()
                            supabase.table("skills").delete().eq("id", sk["id"]).execute()
                            st.session_state[f"confirm_del_sk_{sk['id']}"] = False
                            st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_del_sk_{sk['id']}"):
                            st.session_state[f"confirm_del_sk_{sk['id']}"] = False
                            st.rerun()

                st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── ADD / EDIT SKILL ──
    else:
        editing_sk = st.session_state.get("editing_skill", None)
        editing_sk_id = editing_sk["id"] if editing_sk else None
        st.markdown(f"#### {'✏️ Editing Skill' if editing_sk else 'What do you want to master?'}")

        with st.form("skill_form"):
            sk_name  = st.text_input("Skill Name", value=editing_sk["skill_name"] if editing_sk else "", placeholder="e.g. Japanese, Boxing, Python, Guitar...")
            cats     = ["language", "sport", "tech", "art", "music", "other"]
            sk_cat   = st.selectbox("Category", cats, index=cats.index(editing_sk["category"]) if editing_sk and editing_sk.get("category") in cats else 0)
            levels   = ["beginner", "intermediate", "advanced", "master"]
            sk_level = st.selectbox("Target Level", levels, index=levels.index(editing_sk["target_level"]) if editing_sk and editing_sk.get("target_level") in levels else 0)
            sk_hours = st.number_input("Target Hours to reach that level", min_value=1, step=10, value=int(editing_sk.get("target_hours", 100)) if editing_sk else 100)
            sk_why   = st.text_area("Why do you want this skill?", value=editing_sk.get("motivation","") if editing_sk else "", height=80, placeholder="What drives you to learn this?")
            submitted = st.form_submit_button("💾 Update Skill" if editing_sk else "💾 Add Skill")

        if submitted:
            if not sk_name.strip():
                st.error("Please enter a skill name.")
            else:
                try:
                    record = {
                        "skill_name": sk_name.strip(),
                        "category": sk_cat,
                        "target_level": sk_level,
                        "target_hours": sk_hours,
                        "motivation": sk_why.strip(),
                    }
                    if editing_sk_id:
                        supabase.table("skills").update(record).eq("id", editing_sk_id).execute()
                        st.success("✅ Skill updated.")
                        st.session_state["editing_skill"] = None
                    else:
                        record["date_started"] = str(date.today())
                        supabase.table("skills").insert(record).execute()
                        st.success(f"✅ '{sk_name}' added. Start logging sessions!")
                    st.session_state["skill_view"] = "📊 Active Skills"
                    st.session_state["skill_radio_counter"] += 1
                    st.rerun()
                except Exception as ex:
                    st.error(f"Error: {ex}")

        if editing_sk_id and st.button("Cancel", key="cancel_skill_form"):
            st.session_state["editing_skill"] = None
            st.session_state["skill_view"] = "📊 Active Skills"
            st.session_state["skill_radio_counter"] += 1
            st.rerun()

# ============================================
# SETTINGS PAGE
# ============================================
elif page == "⚙️  Settings":
    st.markdown('<div class="section-header">⚙️ Settings</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-sub">wallpaper · display · preferences</div>', unsafe_allow_html=True)

    st.markdown("#### 🖼 Wallpaper")

    _wp_url_cur = get_setting("wallpaper_url")
    _wp_b64_legacy = get_setting("wallpaper_b64") if not _wp_url_cur else None
    if _wp_url_cur or _wp_b64_legacy:
        st.markdown('<span class="badge badge-green">✅ Wallpaper active</span>', unsafe_allow_html=True)
        if _wp_b64_legacy:
            st.caption("ℹ️ This wallpaper is stored the old way (base64 in the database). "
                       "Re-upload it once to move it to fast Storage.")
        st.markdown("<br>", unsafe_allow_html=True)

    uploaded = st.file_uploader("Upload a wallpaper image", type=["jpg", "jpeg", "png", "webp"], label_visibility="collapsed")
    st.caption("Any photo works — it's compressed, uploaded to Supabase Storage, and only a URL is saved. Fast loads, tiny database.")

    current_opacity = get_overlay_opacity()
    opacity = st.slider("Overlay darkness (how much to dim the wallpaper)", 0.2, 0.9, current_opacity, step=0.05)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save Wallpaper", type="primary"):
            _step = "start"
            try:
                if uploaded:
                    _step = "reading image"
                    img_bytes = uploaded.read()
                    _step = "compressing image"
                    try:
                        from PIL import Image
                        import io as _io
                        _im = Image.open(_io.BytesIO(img_bytes))
                        if _im.mode not in ("RGB", "L"):
                            _im = _im.convert("RGB")
                        _max_w = 1920
                        if _im.width > _max_w:
                            _ratio = _max_w / _im.width
                            _im = _im.resize((_max_w, int(_im.height * _ratio)))
                        _out = _io.BytesIO()
                        _im.save(_out, format="JPEG", quality=82, optimize=True)
                        img_bytes = _out.getvalue()
                    except Exception as _pil_ex:
                        st.caption(f"(compression skipped: {_pil_ex})")

                    _step = f"uploading to Storage ({len(img_bytes):,} bytes)"
                    _fname = "wallpaper.jpg"
                    supabase.storage.from_(WALLPAPER_BUCKET).upload(
                        _fname, img_bytes,
                        {"content-type": "image/jpeg", "upsert": "true"},
                    )
                    _step = "getting public URL"
                    _pub = supabase.storage.from_(WALLPAPER_BUCKET).get_public_url(_fname)
                    # cache-bust so the browser fetches the new image, not yesterday's
                    _pub = f"{_pub.split('?')[0]}?v={int(__import__('time').time())}"

                    _step = "saving wallpaper URL to settings"
                    supabase.table("settings").upsert(
                        {"key": "wallpaper_url", "value": _pub}, on_conflict="key"
                    ).execute()
                    # remove the old heavy base64 row if it exists
                    try:
                        supabase.table("settings").delete().eq("key", "wallpaper_b64").execute()
                    except Exception:
                        pass

                _step = "saving overlay opacity"
                supabase.table("settings").upsert(
                    {"key": "overlay_opacity", "value": str(opacity)}, on_conflict="key"
                ).execute()

                get_setting.clear()  # refresh cached settings immediately
                st.toast("✅ Wallpaper saved.")
                st.rerun()
            except Exception as ex:
                st.error(f"❌ Failed while: **{_step}**")
                _msg = getattr(ex, "message", None) or str(ex)
                _code = getattr(ex, "code", None)
                _hint = getattr(ex, "hint", None)
                _details = getattr(ex, "details", None)
                st.code(f"message: {_msg}\ncode: {_code}\nhint: {_hint}\ndetails: {_details}")
    with col2:
        if (_wp_url_cur or _wp_b64_legacy) and st.button("🗑️ Remove Wallpaper"):
            try:
                supabase.table("settings").delete().in_("key", ["wallpaper_url", "wallpaper_b64"]).execute()
                try:
                    supabase.storage.from_(WALLPAPER_BUCKET).remove(["wallpaper.jpg"])
                except Exception:
                    pass
                get_setting.clear()
                st.toast("🗑️ Wallpaper removed.")
                st.rerun()
            except Exception as ex:
                st.error(f"Error: {ex}")

    if _wp_url_cur or _wp_b64_legacy:
        _prev_src = _wp_url_cur or f"data:image/jpeg;base64,{_wp_b64_legacy}"
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("**Current wallpaper preview:**")
        st.markdown(f'<div style="width:100%;height:120px;background-image:url(\'{_prev_src}\');background-size:cover;background-position:center;border-radius:8px;border:1px solid #2a2a4a;"></div>', unsafe_allow_html=True)

    # ============================================================
    # EXPORT ALL DATA
    # ============================================================
    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-sub" style="letter-spacing:1px;">EXPORT YOUR ARCHIVE</div>', unsafe_allow_html=True)
    st.markdown('<div style="color:#aaa;font-size:0.85rem;margin-bottom:0.8rem;">Download all your data in your preferred format. Run this before using the Danger Zone below.</div>', unsafe_allow_html=True)

    # Format picker — persists for the whole session
    if "export_format" not in st.session_state:
        st.session_state["export_format"] = "Excel (.xlsx)"
    _fmt_opts = ["Excel (.xlsx)", "CSV (.zip)", "PDF", "Word (.docx)"]
    st.session_state["export_format"] = st.radio(
        "Download format", _fmt_opts,
        index=_fmt_opts.index(st.session_state["export_format"]),
        horizontal=True, label_visibility="collapsed", key="export_format_radio"
    )
    _chosen_fmt = st.session_state["export_format"]

    EXPORT_TABLES = [
        "daily_logs", "reading_sessions", "books", "life_events",
        "purchases", "wishes", "goals", "goal_sessions", "outcomes",
        "skills", "skill_sessions", "win_failure_archive",
    ]

    def _fetch_all_tables():
        data = {}
        errors = []
        for t in EXPORT_TABLES:
            try:
                data[t] = supabase.table(t).select("*").execute().data or []
            except Exception as ex:
                data[t] = []
                errors.append(f"{t}: {ex}")
        return data, errors

    def _list_val(v):
        if isinstance(v, list):
            return " | ".join(
                (f"{x.get('type','')} {x.get('label','')}".strip() if isinstance(x, dict) else str(x))
                for x in v
            )
        return str(v) if v is not None else ""

    def _generate_xlsx(tables):
        import io
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = Workbook()
        wb.remove(wb.active)
        for tname, rows in tables.items():
            if not rows: continue
            ws = wb.create_sheet(tname[:31])
            ws.sheet_view.showGridLines = False
            headers = list(rows[0].keys())
            for ci, h in enumerate(headers, 1):
                c = ws.cell(1, ci, h)
                c.font = Font(bold=True, color="A78BFA", name="Arial", size=10)
                c.fill = PatternFill("solid", start_color="1A1A3A")
                c.alignment = Alignment(horizontal="center")
            for ri, row in enumerate(rows, 2):
                for ci, h in enumerate(headers, 1):
                    v = _list_val(row.get(h, ""))
                    c = ws.cell(ri, ci, v)
                    c.font = Font(color="E8E0FF", name="Arial", size=9)
                    c.fill = PatternFill("solid", start_color="0D0D22" if ri%2==0 else "12122A")
                    c.alignment = Alignment(wrap_text=True)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return buf.read()

    def _generate_csv(tables):
        import io, csv, zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for tname, rows in tables.items():
                if not rows: continue
                sbuf = io.StringIO()
                headers = list(rows[0].keys())
                w = csv.DictWriter(sbuf, fieldnames=headers, extrasaction="ignore")
                w.writeheader()
                for row in rows:
                    w.writerow({h: _list_val(row.get(h,"")) for h in headers})
                zf.writestr(f"life_archive/{tname}.csv", sbuf.getvalue().encode("utf-8-sig"))
            zf.writestr("life_archive/README.txt", "Life Archive Export\n\n" + "\n".join(f"  {t}.csv" for t in tables))
        buf.seek(0)
        return buf.read()

    def _generate_pdf(tables):
        import io
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        C_BG = colors.HexColor("#0D0D22"); C_CARD = colors.HexColor("#1A1A3A")
        C_PURP = colors.HexColor("#A78BFA"); C_WHITE = colors.HexColor("#E8E0FF")
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        W = letter[0] - 1.5*inch
        title_s = ParagraphStyle("T", fontName="Helvetica-Bold", fontSize=22, textColor=C_WHITE, spaceAfter=4)
        h2_s    = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=13, textColor=C_PURP, spaceBefore=16, spaceAfter=6)
        story = [Paragraph("Life Archive — Full Export", title_s), Spacer(1, 12)]
        for tname, rows in tables.items():
            if not rows: continue
            story.append(Paragraph(tname.replace("_"," ").title(), h2_s))
            headers = list(rows[0].keys())
            col_w = [W / len(headers)] * len(headers)
            cell_s = ParagraphStyle("td", fontName="Helvetica", fontSize=7, textColor=C_WHITE, leading=9)
            hdr_s  = ParagraphStyle("th", fontName="Helvetica-Bold", fontSize=7, textColor=C_BG)
            data = [[Paragraph(h, hdr_s) for h in headers]]
            for row in rows[:200]:  # cap at 200 rows per table for PDF size
                data.append([Paragraph(_list_val(row.get(h,""))[:80], cell_s) for h in headers])
            t = Table(data, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0), C_PURP),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_CARD, colors.HexColor("#12122A")]),
                ("GRID",(0,0),(-1,-1),0.5, colors.HexColor("#2A2A4A")),
                ("VALIGN",(0,0),(-1,-1),"TOP"),
                ("TOPPADDING",(0,0),(-1,-1),4),
                ("BOTTOMPADDING",(0,0),(-1,-1),4),
            ]))
            story.append(t); story.append(Spacer(1, 8))
        doc.build(story); buf.seek(0)
        return buf.read()

    def _generate_docx(tables):
        import io
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        def rgb(h): h=h.lstrip("#"); return RGBColor(int(h[:2],16),int(h[2:4],16),int(h[4:],16))
        def cell_shd(cell, color):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
            shd.set(qn("w:fill"),color.lstrip("#")); tcPr.append(shd)
        doc = Document()
        for s in doc.sections:
            s.left_margin=s.right_margin=Inches(0.85)
            s.top_margin=s.bottom_margin=Inches(0.85)
        p=doc.add_paragraph(); r=p.add_run("Life Archive — Full Export")
        r.bold=True; r.font.size=Pt(22); r.font.color.rgb=rgb("A78BFA"); r.font.name="Arial"
        for tname, rows in tables.items():
            if not rows: continue
            h2=doc.add_paragraph(); rh=h2.add_run(tname.replace("_"," ").title())
            rh.bold=True; rh.font.size=Pt(13); rh.font.color.rgb=rgb("A78BFA"); rh.font.name="Arial"
            headers=list(rows[0].keys())
            t=doc.add_table(rows=1+min(len(rows),200), cols=len(headers))
            for ci,h in enumerate(headers):
                c=t.rows[0].cells[ci]; cell_shd(c,"1A1A3A")
                p2=c.paragraphs[0]; r2=p2.add_run(h)
                r2.bold=True; r2.font.size=Pt(8); r2.font.color.rgb=rgb("A78BFA"); r2.font.name="Arial"
            for ri,row in enumerate(rows[:200]):
                for ci,h in enumerate(headers):
                    c=t.rows[ri+1].cells[ci]
                    cell_shd(c,"0D0D22" if ri%2==0 else "12122A")
                    p3=c.paragraphs[0]; r3=p3.add_run(_list_val(row.get(h,""))[:100])
                    r3.font.size=Pt(8); r3.font.color.rgb=rgb("E8E0FF"); r3.font.name="Arial"
        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf.read()

    _fmt_map = {
        "Excel (.xlsx)": ("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", _generate_xlsx),
        "CSV (.zip)":    ("zip",  "application/zip", _generate_csv),
        "PDF":           ("pdf",  "application/pdf", _generate_pdf),
        "Word (.docx)":  ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _generate_docx),
    }

    if st.button("📦 Generate full archive export", use_container_width=False):
        with st.spinner(f"Fetching all data and generating {_chosen_fmt}..."):
            _tables, _errors = _fetch_all_tables()
            _ext, _mime, _gen_fn = _fmt_map[_chosen_fmt]
            st.session_state["export_pkg"] = {
                "bytes": _gen_fn(_tables),
                "ext": _ext, "mime": _mime, "fmt": _chosen_fmt,
                "rows": sum(len(v) for v in _tables.values()),
                "errors": _errors,
            }

    _pkg = st.session_state.get("export_pkg")
    if _pkg:
        if _pkg["errors"]:
            st.warning(f"Some tables had issues: {'; '.join(_pkg['errors'])}")
        st.download_button(
            label=f"⬇️ Download {_pkg['fmt']} ({_pkg['rows']} rows)",
            data=_pkg["bytes"],
            file_name=f"life_archive_full_{date.today()}.{_pkg['ext']}",
            mime=_pkg["mime"],
            use_container_width=True,
        )

    # ============================================================
    # DANGER ZONE — wipe all data
    # ============================================================
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # Every table the app writes to. Order doesn't matter — we clear all of them.
    ALL_DATA_TABLES = [
        "daily_logs", "reading_sessions", "books", "life_events",
        "purchases", "wishes", "goals", "goal_sessions", "outcomes",
        "skills", "skill_sessions", "win_failure_archive",
    ]
    CONFIRM_PHRASE = "DELETE EVERYTHING"

    with st.expander("⚠️  Danger Zone — delete all data"):
        st.markdown(
            '<div style="color:#f87171;font-size:0.9rem;">'
            "This permanently erases <b>every entry</b> in your archive — daily logs, "
            "reading, life events, purchases, wishes, goals, outcomes, skills, and the "
            "win/failure archive. This <b>cannot be undone</b> and there are no backups "
            "on the free plan."
            "</div>",
            unsafe_allow_html=True,
        )
        st.markdown("<br>", unsafe_allow_html=True)

        also_wallpaper = st.checkbox("Also reset appearance (remove wallpaper + opacity)", value=False)
        understood = st.checkbox("I understand this is permanent and irreversible.", value=False)
        typed = st.text_input(
            f'Type "{CONFIRM_PHRASE}" to confirm',
            key="danger_confirm",
            placeholder=CONFIRM_PHRASE,
        )

        ready = understood and typed.strip() == CONFIRM_PHRASE
        if st.button("🗑️ Permanently delete all data", type="primary", disabled=not ready):
            tables_to_clear = list(ALL_DATA_TABLES)
            if also_wallpaper:
                tables_to_clear.append("settings")

            results = []
            total_deleted = 0
            # settings is keyed by "key"; every other table by "id"
            id_col_for = {"settings": "key"}
            with st.spinner("Wiping your archive..."):
                for t in tables_to_clear:
                    idcol = id_col_for.get(t, "id")
                    try:
                        rows = supabase.table(t).select(idcol).execute().data or []
                        ids = [r[idcol] for r in rows]
                        # delete in batches so huge tables don't blow the request size limit
                        for _b in range(0, len(ids), 200):
                            supabase.table(t).delete().in_(idcol, ids[_b:_b+200]).execute()
                        results.append((t, len(ids), None))
                        total_deleted += len(ids)
                    except Exception as ex:
                        results.append((t, 0, str(ex)))

            failed = [r for r in results if r[2]]
            if failed:
                st.error(f"Deleted {total_deleted} rows, but some tables failed:")
                for t, _, err in failed:
                    st.markdown(f"&nbsp; ✕ `{t}` — {err}")
            else:
                st.success(f"✅ Done. Permanently deleted {total_deleted} rows across {len(tables_to_clear)} tables.")

            # Clear any cached lists in the current session so the UI reflects the wipe.
            for k in ("accomplishments", "media_list", "workouts"):
                if k in st.session_state:
                    st.session_state[k] = []

        if not ready:
            st.caption("The delete button stays locked until both boxes are checked and the phrase matches exactly.")
